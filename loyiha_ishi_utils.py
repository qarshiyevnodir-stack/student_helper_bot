"""
loyiha_ishi_utils.py — Yangi shablon asosida Loyiha Ishi yaratuvchi modul.

Shablon sahifalari (LOYIHAISHI.docx dizaynidan ilhomlangan):
  Sahifa 1: Muqova (O'zbekiston gerbi, "LOYIHA ISHI", bajardi)
  Sahifa 2: Kirish sahifasi (sarlavha + rasm + matn)
  Sahifa 3+: Almashib keluvchi 4 ta shablon (A, B, C, D)

Sahifa tanlovlari: 5, 10, 15

OPTIMIZATSIYA: Barcha GPT so'rovlari bitta mega-so'rovga birlashtirilgan.
  Eski: 10-20 ta alohida so'rov
  Yangi: 3 ta so'rov (kalit so'z + mega-kontent + xulosa bloki)
"""

import os
import re
import json
import logging
import tempfile
import requests
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
client = OpenAI()

ASSETS_DIR = os.path.join(os.path.dirname(__file__), "assets")
GERB_PATH  = os.path.join(ASSETS_DIR, "gerb.jpeg")
KITOB_PATH = os.path.join(ASSETS_DIR, "kitob.png")

GOLD_COLOR = (200, 160, 0)   # Sariq/oltin rang


# ─────────────────────────────────────────────
# Yordamchi funksiyalar
# ─────────────────────────────────────────────

def strip_markdown(text: str) -> str:
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
    text = re.sub(r'_{1,3}(.*?)_{1,3}', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'^[-\*]{3,}\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def gpt_generate(prompt: str, system: str = "Siz foydali yordamchisiz.") -> str:
    try:
        resp = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system},
                {"role": "user",   "content": prompt},
            ],
        )
        return strip_markdown(resp.choices[0].message.content.strip())
    except Exception as e:
        logging.error(f"GPT xatolik: {e}")
        return f"Kontent yaratishda xatolik: {e}"


def set_font(run, size=14, bold=False, italic=False,
             name='Times New Roman', color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc_or_cell, text='', alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  size=14, bold=False, italic=False,
                  space_before=0, space_after=6, color=None):
    p = doc_or_cell.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_page_border(document):
    """Barcha sahifalarga qora chegara qo'shadi."""
    for section in document.sections:
        sectPr = section._sectPr
        for old in sectPr.findall(qn('w:pgBorders')):
            sectPr.remove(old)
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for edge in ('top', 'left', 'bottom', 'right'):
            border_el = OxmlElement(f'w:{edge}')
            border_el.set(qn('w:val'),   'single')
            border_el.set(qn('w:sz'),    '18')
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), '000000')
            pgBorders.append(border_el)
        sectPr.append(pgBorders)


def remove_cell_borders(cell):
    """Jadval katakchasidan chegara olib tashlanadi."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_gold_line(doc, width=30):
    """Sariq/oltin ajratuvchi chiziq qo'shadi."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * width)
    set_font(run, size=10, color=GOLD_COLOR)
    return p


# ─────────────────────────────────────────────
# Rasm yuklash
# ─────────────────────────────────────────────

def fetch_topic_images(keyword: str, count: int = 10) -> list:
    """Mavzuga oid rasmlarni Unsplash dan yuklab, yo'llar ro'yxatini qaytaradi."""
    images = []
    logging.info(f"Rasm qidirish kalit so'zi: {keyword}")

    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        url = f"https://source.unsplash.com/featured/1200x800/?{requests.utils.quote(keyword)}"
        for i in range(count):
            r = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
            if r.status_code == 200 and r.headers.get('content-type', '').startswith('image'):
                tmp = tempfile.NamedTemporaryFile(suffix='.jpg', delete=False)
                tmp.write(r.content)
                tmp.close()
                images.append(tmp.name)
    except Exception as e:
        logging.warning(f"Unsplash xatolik: {e}")

    # Yetarli rasm bo'lmasa, picsum.photos dan olamiz
    while len(images) < count:
        try:
            r = requests.get(
                f"https://picsum.photos/1200/800?random={len(images) + 100}",
                timeout=10, allow_redirects=True
            )
            if r.status_code == 200:
                tmp = tempfile.NamedTemporaryFile(suffix='.jpg', delete=False)
                tmp.write(r.content)
                tmp.close()
                images.append(tmp.name)
        except Exception:
            break

    return images


# ─────────────────────────────────────────────
# MEGA-SO'ROV: Barcha kontentni bitta so'rovda
# ─────────────────────────────────────────────

def generate_all_content(topic: str, language: str, section_count: int) -> dict:
    """
    OPTIMIZATSIYA: Barcha loyiha ishi kontentini BITTA GPT so'rovida oladi.

    Qaytaradi:
      {
        "keyword": "...",          # Unsplash uchun inglizcha kalit so'z
        "sections": [
          {"title": "...", "text": "..."},
          ...
        ],
        "conclusion": "...",       # Xulosa matni
        "table_data": [            # 4 ta tushuncha
          {"term": "...", "definition": "..."},
          ...
        ],
        "references": ["...", "...", "..."]  # 3 ta adabiyot
      }
    """
    lang_map = {
        'uz': "o'zbek", 'ru': "rus", 'en': "ingliz",
        'ko': "kores",  'zh': "xitoy", 'de': "nemis"
    }
    lang_name = lang_map.get(language, "o'zbek")

    system_msg = (
        f"Siz {lang_name} tilida akademik loyiha ishi yozuvchi mutaxassississiz. "
        f"Faqat sof matn yozing, markdown belgilari ishlatmang. "
        f"Javobingiz to'liq JSON formatida bo'lsin."
    )

    prompt = f"""'{topic}' mavzusida loyiha ishi uchun quyidagi JSON strukturasini to'ldiring.
Til: {lang_name}. Barcha matnlar {lang_name} tilida bo'lsin.

{{
  "keyword": "mavzu uchun eng mos 1-2 ta inglizcha kalit so'z (Unsplash qidiruvi uchun)",
  "sections": [
    {{"title": "1-bo'lim nomi", "text": "150-200 so'zlik akademik matn, paragraflar bilan"}},
    {{"title": "2-bo'lim nomi", "text": "150-200 so'zlik akademik matn, paragraflar bilan"}},
    ... ({section_count} ta bo'lim, birinchisi 'Kirish' bo'lsin)
  ],
  "conclusion": "120 so'zlik xulosa va takliflar matni, asosiy natijalar va amaliy tavsiyalar",
  "table_data": [
    {{"term": "1-tushuncha", "definition": "qisqacha ta'rif"}},
    {{"term": "2-tushuncha", "definition": "qisqacha ta'rif"}},
    {{"term": "3-tushuncha", "definition": "qisqacha ta'rif"}},
    {{"term": "4-tushuncha", "definition": "qisqacha ta'rif"}}
  ],
  "references": [
    "1. Birinchi adabiyot manbasi",
    "2. Ikkinchi adabiyot manbasi",
    "3. Uchinchi adabiyot manbasi"
  ]
}}

Faqat JSON qaytaring, boshqa hech narsa yo'q. {section_count} ta bo'lim bo'lsin."""

    try:
        resp = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user",   "content": prompt},
            ],
            response_format={"type": "json_object"},
        )
        raw = resp.choices[0].message.content.strip()
        data = json.loads(raw)
        logging.info(f"Mega-so'rov muvaffaqiyatli: {len(data.get('sections', []))} bo'lim")
        return data
    except Exception as e:
        logging.error(f"Mega-so'rov xatolik: {e}")
        # Fallback: bo'sh struktura
        return {
            "keyword": topic,
            "sections": [{"title": f"Bo'lim {i+1}", "text": "Matn yaratishda xatolik yuz berdi."} for i in range(section_count)],
            "conclusion": "Xulosa yaratishda xatolik yuz berdi.",
            "table_data": [{"term": "-", "definition": "-"} for _ in range(4)],
            "references": ["1. -", "2. -", "3. -"]
        }


# ─────────────────────────────────────────────
# 1-SAHIFA: Muqova
# ─────────────────────────────────────────────

def create_cover_page(doc, university_info, subject_name, topic,
                      name_surname, teacher_name):
    """O'zbekiston gerbi bilan rasmiy muqova sahifasi."""
    # Gerb
    if os.path.exists(GERB_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        p.add_run().add_picture(GERB_PATH, width=Inches(2.5))

    # Muassasa nomi
    univ_text = university_info.strip().upper() if university_info and university_info.strip() \
        else "TA'LIM MUASSASA NOMI"
    add_paragraph(doc, univ_text,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  size=14, bold=True, space_before=4, space_after=2)

    # Fan nomi
    subj_text = subject_name.strip().upper() if subject_name and subject_name.strip() \
        else "TANLANGAN FANIDAN"
    add_paragraph(doc, subj_text,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  size=14, bold=True, space_before=0, space_after=12)

    # LOYIHA ISHI (katta sarlavha)
    add_paragraph(doc, "LOYIHA ISHI",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  size=36, bold=True, space_before=12, space_after=8)

    # Mavzu
    if topic and topic.strip():
        add_paragraph(doc, f"Mavzu: {topic.strip()}",
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      size=14, bold=True, space_before=4, space_after=16)

    # Bajardi
    p_b = doc.add_paragraph()
    p_b.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_b.paragraph_format.space_before = Pt(4)
    p_b.paragraph_format.space_after  = Pt(2)
    r1 = p_b.add_run("Bajardi: ")
    set_font(r1, size=14, bold=True)
    r2 = p_b.add_run(name_surname.strip() if name_surname else "")
    set_font(r2, size=14, bold=True)

    # O'qituvchi
    p_q = doc.add_paragraph()
    p_q.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_q.paragraph_format.space_before = Pt(2)
    p_q.paragraph_format.space_after  = Pt(20)
    r3 = p_q.add_run("Qabul qildi: ")
    set_font(r3, size=14, bold=True)
    r4 = p_q.add_run(teacher_name.strip() if teacher_name and teacher_name.strip() else "")
    set_font(r4, size=14)

    # Kitob rasmi
    if os.path.exists(KITOB_PATH):
        p_k = doc.add_paragraph()
        p_k.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_k.paragraph_format.space_before = Pt(0)
        p_k.paragraph_format.space_after  = Pt(0)
        p_k.add_run().add_picture(KITOB_PATH, width=Inches(1.8))

    doc.add_page_break()


# ─────────────────────────────────────────────
# 2-SAHIFA: Kirish (sariq dizayn)
# ─────────────────────────────────────────────

def create_intro_page(doc, topic, intro_text, img_path=None,
                      author='', extra_info=''):
    """
    Kirish sahifasi: katta sarlavha + rasm + sariq rangdagi matn bloklari.
    """
    # Katta sarlavha
    add_paragraph(doc, topic,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  size=22, bold=True, space_before=6, space_after=12)

    # Rasm
    if img_path and os.path.exists(img_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(10)
            p.add_run().add_picture(img_path, width=Inches(5.5))
        except Exception as e:
            logging.warning(f"Kirish rasmi: {e}")

    # Kirish matni
    if intro_text:
        for para in intro_text.split('\n'):
            para = para.strip()
            if para:
                add_paragraph(doc, para,
                              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                              size=17, space_before=2, space_after=4)

    doc.add_page_break()


# ─────────────────────────────────────────────
# CONTENT SAHIFALAR: A, B, C, D shablonlari
# ─────────────────────────────────────────────

def create_content_page_A(doc, heading, text_content, img_path=None):
    """Shablon A: Sarlavha yuqorida, chap matn + o'ng rasm."""
    add_paragraph(doc, heading,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  size=20, bold=True, space_before=6, space_after=6)
    add_gold_line(doc, 25)

    table = doc.add_table(rows=1, cols=2)
    remove_cell_borders(table.cell(0, 0))
    remove_cell_borders(table.cell(0, 1))

    left_cell = table.cell(0, 0)
    left_cell.width = Inches(3.6)
    for para in text_content.split('\n'):
        para = para.strip()
        if not para:
            continue
        p = left_cell.add_paragraph(para)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(5)
        for run in p.runs:
            set_font(run, size=16)

    right_cell = table.cell(0, 1)
    right_cell.width = Inches(2.4)
    if img_path and os.path.exists(img_path):
        try:
            p = right_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(img_path, width=Inches(2.2))
        except Exception as e:
            logging.warning(f"Content A rasm: {e}")

    doc.add_page_break()


def create_content_page_B(doc, heading, text_content, img_path=None):
    """Shablon B: Yuqorida to'liq kenglikdagi rasm, pastda sarlavha + matn."""
    if img_path and os.path.exists(img_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(8)
            p.add_run().add_picture(img_path, width=Inches(6.0))
        except Exception as e:
            logging.warning(f"Content B rasm: {e}")

    add_paragraph(doc, heading,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  size=20, bold=True, space_before=6, space_after=2)
    add_gold_line(doc, 25)

    for para in text_content.split('\n'):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph(para)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        for run in p.runs:
            set_font(run, size=16)

    doc.add_page_break()


def create_content_page_C(doc, heading, text_content, img_path=None):
    """Shablon C: Sariq chiziq yuqorida, katta rasm, pastda sarlavha + matn."""
    add_gold_line(doc, 20)

    if img_path and os.path.exists(img_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(8)
            p.add_run().add_picture(img_path, width=Inches(6.0))
        except Exception as e:
            logging.warning(f"Content C rasm: {e}")

    add_paragraph(doc, heading,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  size=24, bold=True, space_before=6, space_after=6)

    for para in text_content.split('\n'):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph(para)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        for run in p.runs:
            set_font(run, size=17, bold=True)

    doc.add_page_break()


def create_content_page_D(doc, heading, text_left, text_right=None):
    """Shablon D: Ikki ustunli matn sahifasi."""
    add_paragraph(doc, heading,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  size=20, bold=True, space_before=6, space_after=6)
    add_gold_line(doc, 25)

    table = doc.add_table(rows=1, cols=2)
    remove_cell_borders(table.cell(0, 0))
    remove_cell_borders(table.cell(0, 1))

    left_cell = table.cell(0, 0)
    left_cell.width = Inches(3.0)
    for para in text_left.split('\n'):
        para = para.strip()
        if not para:
            continue
        p = left_cell.add_paragraph(para)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(5)
        for run in p.runs:
            set_font(run, size=16)

    right_cell = table.cell(0, 1)
    right_cell.width = Inches(3.0)
    content_right = text_right or text_left
    for para in content_right.split('\n'):
        para = para.strip()
        if not para:
            continue
        p = right_cell.add_paragraph(para)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(5)
        for run in p.runs:
            set_font(run, size=16)

    doc.add_page_break()


# ─────────────────────────────────────────────
# OXIRGI SAHIFA: Xulosa + jadval + rasm
# ─────────────────────────────────────────────

def create_final_page(doc, conclusion_text, table_data, references, topic_images, img_idx):
    """
    Oxirgi sahifa: Xulosa matni + taqqoslash jadvali + rasm + adabiyotlar.
    Barcha ma'lumotlar mega-so'rovdan keladi (GPT chaqirilmaydi).
    """
    # Sarlavha
    add_paragraph(doc, "Xulosa va takliflar",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  size=18, bold=True, space_before=6, space_after=4)

    add_gold_line(doc, 25)

    # Xulosa matni
    for para in conclusion_text.split('\n'):
        para = para.strip()
        if para:
            add_paragraph(doc, para,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          size=12, space_before=0, space_after=4)

    # Jadval sarlavhasi
    add_paragraph(doc, "Asosiy tushunchalar jadvali",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  size=13, bold=True, space_before=10, space_after=4,
                  color=GOLD_COLOR)

    # Jadval yaratish
    try:
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Tushuncha"
        hdr_cells[1].text = "Ta'rif"
        for cell in hdr_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, size=12, bold=True, color=GOLD_COLOR)

        rows_data = []
        for item in table_data:
            if isinstance(item, dict):
                term = item.get("term", "-")
                defn = item.get("definition", "-")
                rows_data.append((term, defn))
        while len(rows_data) < 4:
            rows_data.append(("-", "-"))

        for i, (term, definition) in enumerate(rows_data[:4]):
            row_cells = table.rows[i + 1].cells
            row_cells[0].text = str(term)
            row_cells[1].text = str(definition)
            for cell in row_cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_font(run, size=11)
    except Exception as e:
        logging.warning(f"Jadval yaratishda xatolik: {e}")

    # Rasm
    img = topic_images[img_idx % len(topic_images)] if topic_images else None
    if img and os.path.exists(img):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            p.add_run().add_picture(img, width=Inches(4.5))
        except Exception as e:
            logging.warning(f"Oxirgi sahifa rasmi: {e}")

    # Adabiyotlar
    add_paragraph(doc, "Foydalanilgan adabiyotlar",
                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  size=13, bold=True, space_before=10, space_after=4)
    for line in references:
        line = str(line).strip()
        if line:
            add_paragraph(doc, line,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          size=11, space_before=0, space_after=3)


# ─────────────────────────────────────────────
# Asosiy generator
# ─────────────────────────────────────────────

def generate_loyiha_ishi(topic, page_count, language,
                         name_surname, university_info,
                         subject_name, teacher_name):
    """
    To'liq loyiha ishi hujjatini yaratadi va BytesIO qaytaradi.

    Sahifalar tuzilishi:
      1-sahifa : Muqova (har doim)
      2-sahifa : Kirish (har doim)
      3+ sahifa: A, B, C, D, A, B, C, D, ... (almashib keladi)
      Oxirgi   : Xulosa + jadval + adabiyotlar

    Sahifa tanlovlari: 5, 10, 15

    OPTIMIZATSIYA: Barcha kontent BITTA mega-so'rovda olinadi.
    """
    lang_map = {
        'uz': "o'zbek", 'ru': "rus", 'en': "ingliz",
        'ko': "kores",  'zh': "xitoy", 'de': "nemis"
    }
    lang_name = lang_map.get(language, "o'zbek")

    # Nechta content sahifa kerak (muqova va kirish hisoblanmaydi)
    content_page_count = max(3, page_count - 2)
    # Mega-so'rov uchun bo'limlar soni: kirish + content sahifalar
    section_count = content_page_count + 1

    logging.info(f"'{topic}' uchun mega-so'rov yuborilmoqda ({section_count} bo'lim)...")

    # ── SO'ROV 1: Barcha kontent bitta so'rovda ──
    all_content = generate_all_content(topic, language, section_count)

    keyword   = all_content.get("keyword", topic)
    sections  = all_content.get("sections", [])
    conclusion = all_content.get("conclusion", "")
    table_data = all_content.get("table_data", [])
    references = all_content.get("references", [])

    # Bo'limlar yetarli bo'lmasa to'ldirish
    while len(sections) < section_count:
        sections.append({"title": f"Bo'lim {len(sections)+1}", "text": "Matn mavjud emas."})

    # ── SO'ROV 2: Rasmlar yuklash (parallel, GPT emas) ──
    img_count = content_page_count + 2
    logging.info(f"'{keyword}' kalit so'zi bilan {img_count} ta rasm yuklanmoqda...")
    topic_images = fetch_topic_images(keyword, count=img_count)
    logging.info(f"Yuklangan rasmlar: {len(topic_images)}")

    # ── Hujjat yaratish ──
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    for section in doc.sections:
        section.page_width    = Inches(8.27)
        section.page_height   = Inches(11.69)
        section.left_margin   = Inches(1.18)
        section.right_margin  = Inches(0.79)
        section.top_margin    = Inches(0.98)
        section.bottom_margin = Inches(0.98)

    add_page_border(doc)

    # ── 1-SAHIFA: Muqova ──
    create_cover_page(
        doc,
        university_info=university_info,
        subject_name=subject_name,
        topic=topic,
        name_surname=name_surname,
        teacher_name=teacher_name
    )

    # ── 2-SAHIFA: Kirish ──
    img_idx = 0
    intro_section = sections[0]
    intro_text = intro_section.get("text", "")
    intro_img = topic_images[img_idx] if img_idx < len(topic_images) else None
    img_idx += 1

    create_intro_page(
        doc,
        topic=topic,
        intro_text=intro_text,
        img_path=intro_img,
    )

    # ── CONTENT SAHIFALAR (A, B, C, D almashib) ──
    template_cycle = ['A', 'B', 'C', 'D']

    for i in range(content_page_count):
        template_type = template_cycle[i % 4]

        # Bo'lim (sections[1] dan boshlanadi, siklik)
        sec_idx = (i + 1) % len(sections)
        section_data = sections[sec_idx]
        section_title = section_data.get("title", f"Bo'lim {i+1}")
        content = section_data.get("text", "")

        # Rasm (siklik)
        img = topic_images[img_idx % len(topic_images)] if topic_images else None
        img_idx += 1

        logging.info(
            f"Content {i+1}/{content_page_count}: "
            f"shablon={template_type}, bo'lim={section_title!r}"
        )

        if template_type == 'A':
            create_content_page_A(doc, section_title, content, img_path=img)

        elif template_type == 'B':
            create_content_page_B(doc, section_title, content, img_path=img)

        elif template_type == 'C':
            create_content_page_C(doc, section_title, content, img_path=img)

        elif template_type == 'D':
            paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
            mid = len(paragraphs) // 2
            left_text  = '\n'.join(paragraphs[:mid]) if mid > 0 else content
            right_text = '\n'.join(paragraphs[mid:]) if mid < len(paragraphs) else content
            create_content_page_D(doc, section_title, left_text, right_text)

    # ── OXIRGI SAHIFA: Xulosa + jadval + adabiyotlar ──
    create_final_page(doc, conclusion, table_data, references, topic_images, img_idx)

    # ── Vaqtinchalik rasmlarni tozalash ──
    for img_path in topic_images:
        try:
            os.unlink(img_path)
        except Exception:
            pass

    # ── BytesIO ga saqlash ──
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
