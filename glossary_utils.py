"""
glossary_utils.py — Glossary (Atamalar lug'ati) generatori
Berilgan mavzu bo'yicha 15-50 ta atama, ta'rif va misollar bilan professional DOCX hujjat yaratadi.
"""

import os
import json
import logging
from io import BytesIO
from openai import OpenAI

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

def get_client():
    return OpenAI()

# ─────────────────────────────────────────────
# Til sozlamalari
# ─────────────────────────────────────────────

LANG_LABELS = {
    "uz": "O'zbek",
    "ru": "Rus",
    "en": "Ingliz",
    "ko": "Kores",
    "zh": "Xitoy",
    "de": "Nemis",
}

LANG_PROMPTS = {
    "uz": "O'zbek tilida yoz.",
    "ru": "Пиши на русском языке.",
    "en": "Write in English.",
    "ko": "한국어로 작성하세요.",
    "zh": "用中文写。",
    "de": "Schreibe auf Deutsch.",
}

# Glossary hajmi sozlamalari
GLOSSARY_SIZES = {
    "small": {"count": 15, "label": "Kichik (15 ta atama)"},
    "medium": {"count": 30, "label": "O'rta (30 ta atama)"},
    "large": {"count": 50, "label": "Katta (50 ta atama)"},
}

# ─────────────────────────────────────────────
# GPT Mega-so'rov
# ─────────────────────────────────────────────

def generate_glossary_content(topic: str, size: str, lang: str) -> dict:
    """Glossary kontentini GPT dan JSON formatida oladi."""
    cfg = GLOSSARY_SIZES.get(size, GLOSSARY_SIZES["small"])
    count = cfg["count"]
    lang_instruction = LANG_PROMPTS.get(lang, LANG_PROMPTS["uz"])

    prompt = f"""Sen akademik darajadagi terminologiya bo'yicha uzmisan. {lang_instruction}

Mavzu: "{topic}"
Talab qilinadigan atamalar soni: aniq {count} ta

Vazifa: Ushbu mavzuga oid eng muhim va asosiy atamalar lug'atini (glossary) yarat.

Quyidagi qat'iy JSON formatida qaytar:
{{
  "title": "Mavzuga mos glossary sarlavhasi (masalan: Iqtisodiyot terminlari lug'ati)",
  "description": "Ushbu glossary haqida 2-3 jumlalik qisqacha kirish so'zi",
  "terms": [
    {{
      "term": "Atama nomi (Alfavit tartibida)",
      "definition": "Atamaning to'liq va tushunarli akademik ta'rifi (kamida 20 so'z)",
      "example": "Atamaning qo'llanilishiga qisqa misol"
    }},
    ... (jami {count} ta shunday ob'ekt bo'lishi SHART)
  ]
}}

MUHIM TALABLAR:
- Atamalar soni albatta {count} ta bo'lsin.
- Atamalar alfavit tartibida joylashsin.
- Ta'riflar ilmiy va aniq bo'lsin.
- Faqat JSON formatida qaytar, hech qanday qo'shimcha matn qo'shma."""

    client = get_client()
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"},
        temperature=0.7,
        max_tokens=8000,
    )

    return json.loads(response.choices[0].message.content)


# ─────────────────────────────────────────────
# Yordamchi funksiyalar
# ─────────────────────────────────────────────

def add_divider(doc: Document, color_hex: str = "1B3A6B"):
    """Ko'k ajratuvchi chiziq qo'shadi."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ─────────────────────────────────────────────
# DOCX yaratish
# ─────────────────────────────────────────────

def build_glossary_docx(
    content: dict,
    topic: str,
    lang: str,
    author: str = "",
    institution: str = ""
) -> BytesIO:
    """Professional Glossary DOCX hujjatini yaratadi."""

    doc = Document()
    
    # Sahifa sozlamalari
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Standart shrift
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # ─── SARLAVHA QISMI ───
    glossary_labels = {
        "uz": "GLOSSARY (ATAMALAR LUG'ATI)",
        "ru": "ГЛОССАРИЙ (СЛОВАРЬ ТЕРМИНОВ)",
        "en": "GLOSSARY OF TERMS",
        "ko": "용어집 (GLOSSARY)",
        "zh": "词汇表 (GLOSSARY)",
        "de": "GLOSSAR (BEGRIFFSLEXIKON)"
    }
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(glossary_labels.get(lang, "GLOSSARY"))
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(27, 58, 107)

    # Asosiy sarlavha
    title = content.get("title", f"{topic} bo'yicha atamalar lug'ati")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(27, 58, 107)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)

    add_divider(doc)

    # Muallif ma'lumotlari
    if author or institution:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if author:
            run = p.add_run(f"{author}\n")
            run.bold = True
        if institution:
            run = p.add_run(institution)
            run.italic = True
        p.paragraph_format.space_after = Pt(12)

    # Kirish so'zi
    description = content.get("description", "")
    if description:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(description)
        run.italic = True
        p.paragraph_format.space_after = Pt(18)

    # ─── ATAMALAR RO'YXATI ───
    terms = content.get("terms", [])
    
    # Alifbo tartibida saralash (agar GPT adashgan bo'lsa)
    terms = sorted(terms, key=lambda x: x.get("term", "").lower())
    
    for i, item in enumerate(terms, 1):
        term = item.get("term", "")
        definition = item.get("definition", "")
        example = item.get("example", "")
        
        if not term or not definition:
            continue
            
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(6)
        
        # Tartib raqami va Atama (Qalin)
        run = p.add_run(f"{i}. {term} ")
        run.bold = True
        run.font.color.rgb = RGBColor(27, 58, 107)
        
        # Chiziqcha
        run = p.add_run("— ")
        
        # Ta'rif
        run = p.add_run(definition)
        
        # Misol (agar bo'lsa)
        if example:
            example_labels = {
                "uz": "Misol:", "ru": "Пример:", "en": "Example:",
                "ko": "예시:", "zh": "例如:", "de": "Beispiel:"
            }
            ex_label = example_labels.get(lang, "Misol:")
            
            p_ex = doc.add_paragraph()
            p_ex.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_ex.paragraph_format.left_indent = Cm(1.25)
            p_ex.paragraph_format.space_after = Pt(12)
            
            run = p_ex.add_run(f"{ex_label} ")
            run.bold = True
            run.italic = True
            run.font.size = Pt(11)
            
            run = p_ex.add_run(example)
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(80, 80, 80)
        else:
            p.paragraph_format.space_after = Pt(12)

    # BytesIO ga saqlash
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# Asosiy funksiya
# ─────────────────────────────────────────────

async def generate_glossary(
    topic: str,
    size: str,
    lang: str,
    author: str = "",
    institution: str = "",
) -> BytesIO:
    """Glossary yaratadi va BytesIO qaytaradi."""
    import asyncio

    # 1. Kontent generatsiya (mega-so'rov)
    content = await asyncio.to_thread(
        generate_glossary_content, topic, size, lang
    )

    # 2. DOCX yaratish
    docx_bytes = await asyncio.to_thread(
        build_glossary_docx,
        content, topic, lang, author, institution
    )

    return docx_bytes
