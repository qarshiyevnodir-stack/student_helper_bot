"""
taqriz_utils.py — Taqriz xizmati
Asar nomi va mazmuni asosida professional taqriz DOCX yaratadi.
Optimizatsiyalangan: 1 ta GPT so'rov, 8-12 soniya.
"""
import asyncio
import logging
from io import BytesIO
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

def get_client():
    return OpenAI()

TAQRIZ_PRICE = 2000

LANG_PROMPTS = {
    "uz": "O'zbek tilida yoz.",
    "ru": "Пиши на русском.",
    "en": "Write in English.",
    "ko": "한국어로 작성하세요.",
    "zh": "用中文写。",
    "de": "Schreibe auf Deutsch.",
}

LANG_LABELS = {
    "uz": "O'zbek", "ru": "Rus", "en": "Ingliz",
    "ko": "Kores", "zh": "Xitoy", "de": "Nemis",
}

TAQRIZ_TYPES = {
    "kurs":    "Kurs ishi",
    "diplom":  "Diplom ishi",
    "maqola":  "Ilmiy maqola",
    "kitob":   "Kitob",
    "referat": "Referat",
}

# Baholash tizimi
BAHO_LABELS = {
    "az'lo":   "A'lo (5)",
    "yaxshi":  "Yaxshi (4)",
    "qoniqarli": "Qoniqarli (3)",
}


def _generate_taqriz(title: str, doc_type: str, author: str,
                     summary: str, lang: str) -> str:
    lang_inst = LANG_PROMPTS.get(lang, LANG_PROMPTS["uz"])
    type_label = TAQRIZ_TYPES.get(doc_type, doc_type)
    client = get_client()

    prompt = (
        f"{lang_inst} Professional akademik taqriz yoz. "
        f"Asar turi: {type_label}. "
        f"Sarlavha: {title}. "
        f"Muallif: {author}. "
        f"{('Qisqa mazmun: ' + summary + '.') if summary else ''} "
        f"Tuzilma (har bo'limni yangi paragrafda yoz): "
        f"1) UMUMIY TAVSIF: Asarning dolzarbligi va maqsadi (2-3 jumla). "
        f"2) ASOSIY QISM TAHLILI: Tuzilma, metodologiya, asosiy g'oyalar (3-4 jumla). "
        f"3) AFZALLIKLARI: 2-3 ta kuchli tomonni sanab o'tish (2-3 jumla). "
        f"4) KAMCHILIKLARI: 1-2 ta yaxshilash mumkin bo'lgan jihatlar (1-2 jumla). "
        f"5) XULOSA VA BAHO: Umumiy baho va tavsiya (2 jumla). "
        f"Jami 350-450 so'z. Rasmiy, ilmiy uslub. "
        f"Faqat taqriz matnini yoz, sarlavha yozma."
    )
    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1600,
        temperature=0.65,
    )
    return resp.choices[0].message.content.strip()


def _build_taqriz_docx(text: str, title: str, doc_type: str,
                        author: str, reviewer: str, lang: str) -> BytesIO:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)

    def add_run(para, txt, bold=False, size=12, color=None, italic=False):
        r = para.add_run(txt)
        r.bold = bold; r.italic = italic
        r.font.name = "Times New Roman"; r.font.size = Pt(size)
        if color: r.font.color.rgb = color
        return r

    # ── Sarlavha ──
    title_map = {
        "uz": "TAQRIZ", "ru": "РЕЦЕНЗИЯ",
        "en": "REVIEW", "ko": "리뷰", "zh": "评论", "de": "REZENSION"
    }
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title_map.get(lang, "TAQRIZ"))
    tr.bold = True; tr.font.name = "Times New Roman"
    tr.font.size = Pt(16); tr.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    tp.paragraph_format.space_after = Pt(6)

    # ── Asar ma'lumotlari ──
    type_label = TAQRIZ_TYPES.get(doc_type, doc_type)
    on_map = {
        "uz": "ustida", "ru": "на", "en": "on", "ko": "에 대한", "zh": "关于", "de": "zu"
    }
    sub_map = {
        "uz": f"{type_label} {on_map.get(lang,'ustida')}",
        "ru": f"{on_map.get(lang,'на')} {type_label.lower()}",
        "en": f"{on_map.get(lang,'on')} {type_label}",
        "ko": f"{type_label} {on_map.get(lang,'에 대한')}",
        "zh": f"{on_map.get(lang,'关于')} {type_label}",
        "de": f"{on_map.get(lang,'zu')} {type_label}",
    }
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(sub_map.get(lang, f"{type_label} ustida"))
    sr.italic = True; sr.font.name = "Times New Roman"
    sr.font.size = Pt(12); sr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    sp.paragraph_format.space_after = Pt(4)

    # ── Asar sarlavhasi ──
    np_ = doc.add_paragraph()
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = np_.add_run(f'"{title}"')
    nr.bold = True; nr.font.name = "Times New Roman"
    nr.font.size = Pt(13); nr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
    np_.paragraph_format.space_after = Pt(4)

    # ── Muallif ──
    author_map = {
        "uz": "Muallif:", "ru": "Автор:", "en": "Author:",
        "ko": "저자:", "zh": "作者:", "de": "Autor:"
    }
    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(ap, f"{author_map.get(lang,'Muallif:')}  ", bold=True, size=11)
    add_run(ap, author, size=11, italic=True)
    ap.paragraph_format.space_after = Pt(14)

    # ── Ajratuvchi chiziq ──
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(14)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), '1A73E8')
    pBdr.append(b); pPr.append(pBdr)

    # ── Taqriz matni ──
    # Bo'limlarni ajratish
    section_keywords = [
        "UMUMIY TAVSIF", "ASOSIY QISM TAHLILI", "AFZALLIKLARI",
        "KAMCHILIKLARI", "XULOSA VA BAHO",
        "ОБЩАЯ ХАРАКТЕРИСТИКА", "АНАЛИЗ ОСНОВНОЙ ЧАСТИ", "ДОСТОИНСТВА",
        "НЕДОСТАТКИ", "ЗАКЛЮЧЕНИЕ",
        "GENERAL DESCRIPTION", "ANALYSIS", "STRENGTHS",
        "WEAKNESSES", "CONCLUSION",
    ]

    for para_text in text.split("\n"):
        para_text = para_text.strip()
        if not para_text:
            continue

        # Bo'lim sarlavhasi ekanligini tekshirish
        is_section = any(kw in para_text.upper() for kw in section_keywords)
        # Raqam bilan boshlangan bo'lim (1) UMUMIY...)
        if para_text and para_text[0].isdigit() and ")" in para_text[:4]:
            is_section = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5

        if is_section:
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(para_text)
            r.bold = True; r.font.name = "Times New Roman"
            r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        else:
            p.paragraph_format.first_line_indent = Cm(1.25)
            r = p.add_run(para_text)
            r.font.name = "Times New Roman"; r.font.size = Pt(12)

    # ── Taqrizchi imzosi ──
    doc.add_paragraph()
    sep2 = doc.add_paragraph()
    sep2.paragraph_format.space_before = Pt(20)
    sep2.paragraph_format.space_after = Pt(10)
    pPr2 = sep2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    b2 = OxmlElement('w:bottom')
    b2.set(qn('w:val'), 'single'); b2.set(qn('w:sz'), '4')
    b2.set(qn('w:space'), '1'); b2.set(qn('w:color'), 'CCCCCC')
    pBdr2.append(b2); pPr2.append(pBdr2)

    reviewer_map = {
        "uz": "Taqrizchi:", "ru": "Рецензент:", "en": "Reviewer:",
        "ko": "검토자:", "zh": "审阅者:", "de": "Gutachter:"
    }
    rp = doc.add_paragraph()
    rp.paragraph_format.space_after = Pt(4)
    add_run(rp, f"{reviewer_map.get(lang,'Taqrizchi:')}  ", bold=True, size=11)
    if reviewer:
        add_run(rp, reviewer, size=11)
    else:
        add_run(rp, "_______________________", size=11,
                color=RGBColor(0x88, 0x88, 0x88))

    # Imzo chizig'i
    ip = doc.add_paragraph()
    ip.paragraph_format.space_after = Pt(4)
    sign_map = {
        "uz": "Imzo:", "ru": "Подпись:", "en": "Signature:",
        "ko": "서명:", "zh": "签名:", "de": "Unterschrift:"
    }
    add_run(ip, f"{sign_map.get(lang,'Imzo:')}  ", bold=True, size=11)
    add_run(ip, "_______________________", size=11,
            color=RGBColor(0x88, 0x88, 0x88))

    # Sana chizig'i
    date_map = {
        "uz": "Sana:", "ru": "Дата:", "en": "Date:",
        "ko": "날짜:", "zh": "日期:", "de": "Datum:"
    }
    dp = doc.add_paragraph()
    dp.paragraph_format.space_after = Pt(4)
    add_run(dp, f"{date_map.get(lang,'Sana:')}  ", bold=True, size=11)
    add_run(dp, "_______________________", size=11,
            color=RGBColor(0x88, 0x88, 0x88))

    # ── Footer ──
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(20)
    fr = fp.add_run("@slidego | t.me/slidego")
    fr.font.name = "Times New Roman"; fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    fr.italic = True

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf


async def generate_taqriz(title: str, doc_type: str, author: str,
                           reviewer: str, lang: str, summary: str = "") -> BytesIO:
    text = await asyncio.to_thread(_generate_taqriz, title, doc_type, author, summary, lang)
    return await asyncio.to_thread(_build_taqriz_docx, text, title, doc_type, author, reviewer, lang)
