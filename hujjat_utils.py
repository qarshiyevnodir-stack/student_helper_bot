"""
hujjat_utils.py — Hujjat & Dizayn xizmatlari
1. Rezyume / CV     — DOCX
2. Motivatsion xat  — DOCX
3. Jadval           — Excel (.xlsx)
4. Kontsept xarita  — PNG (matplotlib)

Optimizatsiyalangan: qisqa prompt, 1 ta so'rov, 5-10 soniya.
"""
import os
import json
import logging
import asyncio
from io import BytesIO
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from jinja2 import Template
import weasyprint

logger = logging.getLogger(__name__)

def get_client():
    return OpenAI()

# ─────────────────────────────────────────────
# Narxlar
# ─────────────────────────────────────────────
HUJJAT_PRICES = {
    "rezyume":    3000,
    "motivatsion": 2000,
    "jadval":     2000,
    "mindmap":    2000,
}

LANG_PROMPTS = {
    "uz": "O'zbek tilida yoz.",
    "ru": "Пиши на русском.",
    "en": "Write in English.",
    "ko": "한국어로 작성하세요.",
    "zh": "用中文写。",
    "de": "Schreibe auf Deutsch.",
    "tr": "Türkçe yaz.",
    "tg": "Ба забони тоҷикӣ нависед.",
    "kaa": "Qaraqalpaq tilinde jaz.",
    "kk": "Қазақ тілінде жазыңыз.",
}
LANG_LABELS = {
    "uz": "O'zbek", "ru": "Rus", "en": "Ingliz",
    "ko": "Kores", "zh": "Xitoy", "de": "Nemis",
}

# ── Til bo'yicha bo'lim sarlavhalari ──
CV_LABELS = {
    "uz": {
        "summary": "Qisqacha ma'lumot", "experience": "Ish tajribasi",
        "education": "Ta'lim", "skills": "Ko'nikmalar", "languages": "Tillar",
        "contact": "Aloqa", "projects": "Loyihalar", "certifications": "Sertifikatlar",
        "interests": "Qiziqishlar"
    },
    "ru": {
        "summary": "Профессиональное резюме", "experience": "Опыт работы",
        "education": "Образование", "skills": "Навыки", "languages": "Языки",
        "contact": "Контакты", "projects": "Проекты", "certifications": "Сертификаты",
        "interests": "Интересы"
    },
    "en": {
        "summary": "Professional Summary", "experience": "Experience",
        "education": "Education", "skills": "Skills", "languages": "Languages",
        "contact": "Contact", "projects": "Projects", "certifications": "Certifications",
        "interests": "Interests"
    },
    "ko": {
        "summary": "자기소개", "experience": "경력",
        "education": "학력", "skills": "기술", "languages": "언어",
        "contact": "연락처", "projects": "프로젝트", "certifications": "자격증",
        "interests": "관심사"
    },
    "zh": {
        "summary": "个人简介", "experience": "工作经历",
        "education": "教育背景", "skills": "技能", "languages": "语言",
        "contact": "联系方式", "projects": "项目", "certifications": "证书",
        "interests": "兴趣爱好"
    },
    "de": {
        "summary": "Berufsprofil", "experience": "Berufserfahrung",
        "education": "Ausbildung", "skills": "Fähigkeiten", "languages": "Sprachen",
        "contact": "Kontakt", "projects": "Projekte", "certifications": "Zertifikate",
        "interests": "Interessen"
    },
    "tr": {
        "summary": "Profesyonel Özet", "experience": "İş Deneyimi",
        "education": "Eğitim", "skills": "Beceriler", "languages": "Diller",
        "contact": "İletişim", "projects": "Projeler", "certifications": "Sertifikalar",
        "interests": "İlgi Alanları"
    },
    "tg": {
        "summary": "Хулосаи касбӣ", "experience": "Таҷрибаи корӣ",
        "education": "Маълумот", "skills": "Малакаҳо", "languages": "Забонҳо",
        "contact": "Тамос", "projects": "Лоиҳаҳо", "certifications": "Сертификатҳо",
        "interests": "Шавқу ҳавасҳо"
    },
    "kaa": {
        "summary": "Қысқаша мағлыўмат", "experience": "Жумыс тәжирийбеси",
        "education": "Билими", "skills": "Көникпелер", "languages": "Тиллер",
        "contact": "Байланыс", "projects": "Жобалар", "certifications": "Сертификатлар",
        "interests": "Қызығыўлар"
    },
    "kk": {
        "summary": "Кәсіби түйін", "experience": "Жұмыс тәжірибесі",
        "education": "Білім", "skills": "Дағдылар", "languages": "Тілдер",
        "contact": "Байланыс", "projects": "Жобалар", "certifications": "Сертификаттар",
        "interests": "Қызығушылықтар"
    },
}

# ═══════════════════════════════════════════════
# 1. REZYUME / CV  (HTML → PDF)
# ═══════════════════════════════════════════════
def _generate_cv_content(name: str, profession: str, lang: str, extra: str = "") -> dict:
    """Professional CV kontent — JSON formatida."""
    lang_inst = LANG_PROMPTS.get(lang, LANG_PROMPTS["uz"])
    client = get_client()
    prompt = (
        f"{lang_inst} Professional CV yoz. "
        f"Ism: {name}. Kasb: {profession}. "
        f"{('Qoshimcha: ' + extra) if extra else ''} "
        f"JSON formatida qaytar: "
        f"{{\"summary\": \"3-4 jumlali professional tavsif\", "
        f"\"skills\": [\"skill1\", \"skill2\", \"skill3\", \"skill4\", \"skill5\", \"skill6\", \"skill7\", \"skill8\"], "
        f"\"experience\": ["
        f"  {{\"title\":\"lavozim\",\"company\":\"kompaniya\",\"date\":\"2021 - hozir\",\"bullets\":[\"yutuq 1\",\"yutuq 2\",\"yutuq 3\"]}},"
        f"  {{\"title\":\"lavozim2\",\"company\":\"kompaniya2\",\"date\":\"2018-2021\",\"bullets\":[\"yutuq 1\",\"yutuq 2\"]}}"
        f"], "
        f"\"education\": ["
        f"  {{\"degree\":\"daraja va mutaxassislik\",\"school\":\"universitet nomi\",\"date\":\"2014-2018\",\"description\":\"GPA 3.8, diplom bilan\"}}"
        f"], "
        f"\"projects\": ["
        f"  {{\"name\":\"loyiha nomi\",\"description\":\"qisqa tavsif 1-2 jumla\"}}"
        f"], "
        f"\"certifications\": ["
        f"  {{\"name\":\"sertifikat nomi\",\"organization\":\"tashkilot\",\"date\":\"2023\"}}"
        f"], "
        f"\"languages\": ["
        f"  {{\"name\":\"til nomi\",\"level\":\"daraja\"}}"
        f"], "
        f"\"interests\": [\"qiziqish1\", \"qiziqish2\"], "
        f"\"contact\": {{\"phone\":\"+998 XX XXX XX XX\",\"email\":\"email@example.com\",\"location\":\"Toshkent, O'zbekiston\",\"linkedin\":\"\",\"telegram\":\"\"}} }}"
    )
    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=2500,
        temperature=0.6,
        response_format={"type": "json_object"},
    )
    return json.loads(resp.choices[0].message.content)


def _build_cv_pdf(data: dict, name: str, profession: str, lang: str) -> BytesIO:
    """HTML → PDF rezyume — @slaydtopbot stilida."""
    labels = CV_LABELS.get(lang, CV_LABELS["uz"])

    # HTML shablonini o'qi
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resume_template.html")
    with open(template_path, "r", encoding="utf-8") as f:
        template_str = f.read()

    tmpl = Template(template_str)
    html = tmpl.render(
        lang_code=lang,
        full_name=name.upper(),
        job_title=profession,
        labels=labels,
        summary=data.get("summary", ""),
        experience=data.get("experience", []),
        education=data.get("education", []),
        projects=data.get("projects", []),
        certifications=data.get("certifications", []),
        skills=data.get("skills", []),
        languages=data.get("languages", []),
        interests=data.get("interests", []),
        contact=data.get("contact", {}),
    )

    # HTML → PDF
    pdf_bytes = weasyprint.HTML(string=html, base_url=".").write_pdf()
    buf = BytesIO(pdf_bytes)
    buf.seek(0)
    return buf



def _cv_clean(value) -> str:
    """Bo'sh yoki o'tkazib yuborilgan CV maydonlarini bir xil ko'rinishga keltiradi."""
    if value is None:
        return ""
    value = str(value).strip()
    return "" if value.lower() in {"", "-", "yo'q", "yoq", "none", "n/a"} else value


def _generate_cv_full_content(cv_data: dict) -> dict:
    """Foydalanuvchi bergan ma'lumotni fakt to'qimasdan CV JSONiga normalizatsiya qiladi."""
    lang = cv_data.get("lang", "uz")
    lang_inst = LANG_PROMPTS.get(lang, LANG_PROMPTS["uz"])
    client = get_client()

    raw = {
        "fullname": _cv_clean(cv_data.get("fullname")),
        "email": _cv_clean(cv_data.get("email")),
        "phone": _cv_clean(cv_data.get("phone")),
        "location": _cv_clean(cv_data.get("location")),
        "links": _cv_clean(cv_data.get("links")),
        "title": _cv_clean(cv_data.get("title")),
        "region": _cv_clean(cv_data.get("region")),
        "summary": _cv_clean(cv_data.get("summary")),
        "experience": _cv_clean(cv_data.get("experience")),
        "projects": _cv_clean(cv_data.get("projects")),
        "education": _cv_clean(cv_data.get("education")),
        "certifications": _cv_clean(cv_data.get("certifications")),
        "skills": _cv_clean(cv_data.get("skills")),
        "languages": _cv_clean(cv_data.get("languages")),
    }
    tone = cv_data.get("tone", "professional")

    prompt = f"""{lang_inst}
You are a precise CV editor. Convert ONLY the user's information below into clean CV JSON.

NON-NEGOTIABLE RULES:
- Never invent, assume, embellish, or add any employer, project, date, degree, certificate, language, number, achievement, responsibility, or skill.
- If a source field is empty, the corresponding JSON value must be an empty string or empty list.
- You may correct spelling, split provided text into entries, and make wording concise without changing facts.
- Do not write placeholder values such as 'N/A', 'not specified', 'company', or 'date'.
- Keep the selected tone ({tone}) only in the wording style; do not change factual content.

USER DATA:
Full name: {raw['fullname']}
Email: {raw['email']}
Phone: {raw['phone']}
Location: {raw['location']}
Links: {raw['links']}
Target title: {raw['title']}
Region: {raw['region']}
Professional summary: {raw['summary']}
Experience: {raw['experience']}
Projects: {raw['projects']}
Education: {raw['education']}
Certifications: {raw['certifications']}
Skills: {raw['skills']}
Languages: {raw['languages']}

Return only this JSON object:
{{
  "summary": "",
  "skills": [],
  "experience": [{{"title":"","company":"","date":"","bullets":[]}}],
  "education": [{{"degree":"","school":"","date":"","description":""}}],
  "projects": [{{"name":"","description":""}}],
  "certifications": [{{"name":"","organization":"","date":""}}],
  "languages": [{{"name":"","level":""}}],
  "interests": []
}}"""
    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=2200,
        temperature=0.15,
        response_format={"type": "json_object"},
    )
    data = json.loads(resp.choices[0].message.content)

    # Qo'shimcha himoya: foydalanuvchi bo'sh qoldirgan bo'limlar hech qachon chiqmaydi.
    for key in ("summary", "experience", "projects", "education", "certifications", "skills", "languages"):
        if not raw.get(key):
            data[key] = "" if key == "summary" else []
    data["contact"] = {
        "phone": raw["phone"],
        "email": raw["email"],
        "location": raw["location"],
        "links": raw["links"],
    }
    data["interests"] = []
    return data

def _prepare_cv_for_length(data: dict, length: int) -> dict:
    """1 sahifalik CV uchun faqat eng muhim ma'lumotlarni ixchamlashtiradi."""
    import copy
    prepared = copy.deepcopy(data)
    if length != 1:
        return prepared
    prepared["experience"] = prepared.get("experience", [])[:2]
    for exp in prepared["experience"]:
        exp["bullets"] = exp.get("bullets", [])[:2]
    prepared["projects"] = prepared.get("projects", [])[:1]
    prepared["certifications"] = prepared.get("certifications", [])[:2]
    prepared["skills"] = prepared.get("skills", [])[:10]
    prepared["languages"] = prepared.get("languages", [])[:3]
    return prepared


def _build_cv_full_pdf(data: dict, cv_data: dict) -> BytesIO:
    """Tanlangan dizaynda standart A4 PDF CV yaratadi."""
    lang = cv_data.get("lang", "uz")
    labels = CV_LABELS.get(lang, CV_LABELS["uz"])
    style = cv_data.get("style", "professional")
    length = int(cv_data.get("length", 1))
    template_names = {
        "minimal": "resume_template_minimal.html",
        "professional": "resume_template.html",
        "creative": "resume_template_creative.html",
    }
    template_name = template_names.get(style, "resume_template.html")
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), template_name)
    with open(template_path, "r", encoding="utf-8") as f:
        template_str = f.read()
    data = _prepare_cv_for_length(data, length)

    from jinja2 import Template
    tmpl = Template(template_str)
    
    # Process photo
    photo_data = cv_data.get("photo")
    photo_b64 = ""
    if photo_data:
        import base64
        photo_b64 = f"data:image/jpeg;base64,{base64.b64encode(photo_data).decode('utf-8')}"
        
    # contact dict ni foydalanuvchi ma'lumotlari bilan boyitish
    contact = data.get("contact", {})
    if not contact.get("phone") and cv_data.get("phone"):
        contact["phone"] = cv_data.get("phone", "")
    if not contact.get("email") and cv_data.get("email"):
        contact["email"] = cv_data.get("email", "")
    if not contact.get("location") and cv_data.get("location"):
        contact["location"] = cv_data.get("location", "")
    
    html = tmpl.render(
        lang_code=lang,
        full_name=cv_data.get("fullname", "").upper(),
        job_title=cv_data.get("title", ""),
        labels=labels,
        summary=data.get("summary", ""),
        experience=data.get("experience", []),
        education=data.get("education", []),
        projects=data.get("projects", []),
        certifications=data.get("certifications", []),
        skills=data.get("skills", []),
        languages=data.get("languages", []),
        interests=data.get("interests", []),
        contact=contact,
        photo_b64=photo_b64,
        links=cv_data.get("links", ""),
        region=cv_data.get("region", ""),
        length=length,
    )

    import weasyprint
    from weasyprint import CSS
    # Haqiqiy A4: kontent uzun bo'lsa WeasyPrint avtomatik 2-sahifaga o'tadi.
    a4_css = CSS(string='@page { size: A4; margin: 0; }')
    pdf_bytes = weasyprint.HTML(string=html, base_url=".").write_pdf(stylesheets=[a4_css])

    buf = BytesIO(pdf_bytes)
    buf.seek(0)
    return buf

def _docx_section_title(document, title: str, accent: RGBColor):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = accent


def _build_cv_full_docx(data: dict, cv_data: dict) -> BytesIO:
    """Tahrir qilinadigan, professional DOCX CV yaratadi."""
    data = _prepare_cv_for_length(data, int(cv_data.get("length", 1)))
    lang = cv_data.get("lang", "uz")
    labels = CV_LABELS.get(lang, CV_LABELS["uz"])
    style = cv_data.get("style", "professional")
    accent_values = {
        "minimal": (0x1F, 0x29, 0x37),
        "professional": (0x00, 0x78, 0x6E),
        "creative": (0x6D, 0x28, 0xD9),
    }
    accent = RGBColor(*accent_values.get(style, accent_values["professional"]))

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)

    # Sarlavha va oxirgi qadamda yuborilgan profil rasmi.
    # Rasm yo'q bo'lsa ism butun foydali kenglikdan foydalanadi va keraksiz satrga bo'linmaydi.
    photo = cv_data.get("photo")
    if photo:
        header = document.add_table(rows=1, cols=2)
        header.autofit = False
        header.columns[0].width = Cm(13.2)
        header.columns[1].width = Cm(3.8)
        left_cell, right_cell = header.rows[0].cells
        left_cell.width = Cm(13.2)
        right_cell.width = Cm(3.8)
        p = left_cell.paragraphs[0]
    else:
        right_cell = None
        p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(_cv_clean(cv_data.get("fullname")).upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = accent
    if _cv_clean(cv_data.get("title")):
        role_p = left_cell.add_paragraph() if photo else document.add_paragraph()
        role = role_p.add_run(_cv_clean(cv_data.get("title")))
        role.bold = True
        role.font.name = "Arial"
        role.font.size = Pt(12)

    if photo and right_cell:
        photo_p = right_cell.paragraphs[0]
        photo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        photo_p.add_run().add_picture(BytesIO(photo), width=Cm(3.0), height=Cm(4.0))

    contact_values = [
        _cv_clean(cv_data.get("phone")),
        _cv_clean(cv_data.get("email")),
        _cv_clean(cv_data.get("location")),
        _cv_clean(cv_data.get("links")),
    ]
    contact_text = "  |  ".join(v for v in contact_values if v)
    if contact_text:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        contact_run = p.add_run(contact_text)
        contact_run.font.name = "Arial"
        contact_run.font.size = Pt(9)

    if data.get("summary"):
        _docx_section_title(document, labels["summary"], accent)
        document.add_paragraph(data["summary"])

    if data.get("experience"):
        _docx_section_title(document, labels["experience"], accent)
        for exp in data["experience"]:
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            heading = " | ".join(v for v in [exp.get("title", ""), exp.get("company", "")] if v)
            r = p.add_run(heading)
            r.bold = True
            r.font.name = "Arial"
            if exp.get("date"):
                date = p.add_run(f"    {exp['date']}")
                date.italic = True
                date.font.size = Pt(9)
            for bullet in exp.get("bullets", []):
                bullet_p = document.add_paragraph(style="List Bullet")
                bullet_p.add_run(bullet)

    if data.get("education"):
        _docx_section_title(document, labels["education"], accent)
        for edu in data["education"]:
            p = document.add_paragraph()
            title = " | ".join(v for v in [edu.get("degree", ""), edu.get("school", "")] if v)
            r = p.add_run(title)
            r.bold = True
            if edu.get("date"):
                p.add_run(f"    {edu['date']}")
            if edu.get("description"):
                document.add_paragraph(edu["description"])

    if data.get("projects"):
        _docx_section_title(document, labels["projects"], accent)
        for project in data["projects"]:
            p = document.add_paragraph()
            r = p.add_run(project.get("name", ""))
            r.bold = True
            if project.get("description"):
                p.add_run(f" — {project['description']}")

    if data.get("skills"):
        _docx_section_title(document, labels["skills"], accent)
        document.add_paragraph(" • ".join(data["skills"]))

    if data.get("languages"):
        _docx_section_title(document, labels["languages"], accent)
        language_items = [
            " — ".join(v for v in [item.get("name", ""), item.get("level", "")] if v)
            for item in data["languages"]
        ]
        document.add_paragraph(" • ".join(item for item in language_items if item))

    if data.get("certifications"):
        _docx_section_title(document, labels["certifications"], accent)
        for cert in data["certifications"]:
            item = " | ".join(v for v in [cert.get("name", ""), cert.get("organization", ""), cert.get("date", "")] if v)
            if item:
                document.add_paragraph(item, style="List Bullet")

    buf = BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf


async def generate_cv_full(cv_data: dict) -> BytesIO:
    data = await asyncio.to_thread(_generate_cv_full_content, cv_data)
    return await asyncio.to_thread(_build_cv_full_pdf, data, cv_data)


async def generate_cv_full_docx(cv_data: dict) -> BytesIO:
    data = await asyncio.to_thread(_generate_cv_full_content, cv_data)
    return await asyncio.to_thread(_build_cv_full_docx, data, cv_data)

async def generate_cv(name: str, profession: str, lang: str, extra: str = "") -> BytesIO:
    data = await asyncio.to_thread(_generate_cv_content, name, profession, lang, extra)
    return await asyncio.to_thread(_build_cv_pdf, data, name, profession, lang)


# ═══════════════════════════════════════════════
# 2. MOTIVATSION XAT
# ═══════════════════════════════════════════════
def _generate_motivation_content(name: str, target: str, lang: str, reason: str = "") -> str:
    """Optimizatsiyalangan motivatsion xat — 1 so'rov, to'g'ridan-to'g'ri matn."""
    lang_inst = LANG_PROMPTS.get(lang, LANG_PROMPTS["uz"])
    client = get_client()
    prompt = (
        f"{lang_inst} Professional motivatsion xat yoz. "
        f"Muallif: {name}. Maqsad: {target}. "
        f"{('Sabab: ' + reason) if reason else ''} "
        f"Tuzilma: 1) Murojaat jumlasi, 2) O'zim haqida (2 paragraf), "
        f"3) Nima uchun shu joy (1 paragraf), 4) Xulosa. "
        f"Rasmiy, ishonchli, 350-450 so'z. "
        f"MUHIM: Faqat xat matnini yoz. Imzo yozma — uni alohida qo'shaman."
    )
    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1400,
        temperature=0.65,
    )
    return resp.choices[0].message.content.strip()


def _build_motivation_docx(text: str, name: str, target: str, lang: str) -> BytesIO:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(3.0); sec.bottom_margin = Cm(2.5)

    def add_run(para, text, bold=False, size=12, color=None, italic=False):
        r = para.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.name = "Times New Roman"; r.font.size = Pt(size)
        if color: r.font.color.rgb = color
        return r

    # ── Sarlavha ──
    title_map = {
        "uz": "MOTIVATSION XAT", "ru": "МОТИВАЦИОННОЕ ПИСЬМО",
        "en": "MOTIVATION LETTER", "ko": "자기소개서", "zh": "动机信", "de": "MOTIVATIONSSCHREIBEN"
    }
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title_map.get(lang, "MOTIVATSION XAT"))
    tr.bold = True; tr.font.name = "Times New Roman"
    tr.font.size = Pt(16); tr.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after = Pt(6)

    # ── Maqsad ──
    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = mp.add_run(target)
    mr.italic = True; mr.font.name = "Times New Roman"
    mr.font.size = Pt(12); mr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    mp.paragraph_format.space_after = Pt(18)

    # ── Ajratuvchi ──
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(14)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), '1A73E8')
    pBdr.append(b); pPr.append(pBdr)

    # ── Xat matni ──
    # Imzo jumlalarini matndan tozalash
    sign_keywords = [
        "hurmat bilan", "с уважением", "sincerely", "mit freundlichen",
        "경의를 표하며", "此致", name.lower().split()[0] if name else ""
    ]
    lines = text.split("\n")
    clean_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Imzo qismini o'tkazib yuborish
        low = stripped.lower()
        if any(kw in low for kw in sign_keywords if kw):
            continue
        clean_lines.append(stripped)

    for para_text in clean_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(para_text)
        r.font.name = "Times New Roman"; r.font.size = Pt(12)

    # ── Imzo (bir marta, o'ngda) ──
    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    sign_labels = {
        "uz": "Hurmat bilan,",
        "ru": "С уважением,",
        "en": "Sincerely,",
        "ko": "경의를 표하며,",
        "zh": "此致,",
        "de": "Mit freundlichen Grüßen,"
    }
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sr = sp.add_run(f"{sign_labels.get(lang, 'Hurmat bilan,')}\n{name}")
    sr.font.name = "Times New Roman"; sr.font.size = Pt(12)
    sp.paragraph_format.space_after = Pt(4)

    # ── Footer ──
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(20)
    fr = fp.add_run("@slidego | t.me/slidego")
    fr.font.name = "Times New Roman"; fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    fr.italic = True

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf


async def generate_motivation(name: str, target: str, lang: str, reason: str = "") -> BytesIO:
    text = await asyncio.to_thread(_generate_motivation_content, name, target, lang, reason)
    return await asyncio.to_thread(_build_motivation_docx, text, name, target, lang)


# ═══════════════════════════════════════════════
# 3. JADVAL / DIAGRAMMA (Excel)
# ═══════════════════════════════════════════════
def _generate_table_content(topic: str, lang: str) -> dict:
    """Jadval ma'lumotlarini GPT dan oladi."""
    lang_inst = LANG_PROMPTS.get(lang, LANG_PROMPTS["uz"])
    client = get_client()
    prompt = (
        f"{lang_inst} {topic} mavzusida aniq, haqiqiy jadval ma'lumotlari yarat. "
        f"JSON: {{\"title\": \"jadval sarlavhasi\", "
        f"\"headers\": [\"ustun1\", \"ustun2\", \"ustun3\", \"ustun4\"], "
        f"\"rows\": [[\"qiymat1\", \"qiymat2\", \"qiymat3\", \"qiymat4\"], ...], "
        f"\"summary\": \"qisqa xulosa\"}} "
        f"10-12 qator, 4-5 ustun. Raqamli ma'lumotlar bo'lsin (diagramma uchun)."
    )
    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1800,
        temperature=0.5,
        response_format={"type": "json_object"},
    )
    return json.loads(resp.choices[0].message.content)


def _build_excel(data: dict, topic: str) -> BytesIO:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, Reference
    from openpyxl.chart.series import DataPoint

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = topic[:28]

    title = data.get("title", topic)
    headers = data.get("headers", [])
    rows = data.get("rows", [])
    summary = data.get("summary", "")

    thin = Side(style="thin", color="D0D8E8")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # ── Sarlavha ──
    end_col = chr(64 + max(len(headers), 1))
    ws.merge_cells(f"A1:{end_col}1")
    ws["A1"] = title
    ws["A1"].font = Font(name="Arial", bold=True, size=14, color="1A1A6E")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws["A1"].fill = PatternFill("solid", fgColor="EBF3FE")
    ws.row_dimensions[1].height = 32

    # ── Header qatori ──
    header_fill = PatternFill("solid", fgColor="1A73E8")
    header_font = Font(name="Arial", bold=True, size=11, color="FFFFFF")

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
        ws.column_dimensions[chr(64+col)].width = max(16, len(str(h))+5)
    ws.row_dimensions[2].height = 28

    # ── Ma'lumot qatorlari ──
    for r_idx, row in enumerate(rows, 3):
        fill_color = "F0F6FF" if r_idx % 2 == 0 else "FFFFFF"
        row_fill = PatternFill("solid", fgColor=fill_color)
        for c_idx, val in enumerate(row, 1):
            # Raqamga o'tkazishga harakat
            try:
                val_num = float(str(val).replace(",", ".").replace(" ", ""))
                display_val = val_num
            except (ValueError, TypeError):
                display_val = val
            cell = ws.cell(row=r_idx, column=c_idx, value=display_val)
            cell.font = Font(name="Arial", size=10)
            cell.fill = row_fill
            cell.border = border
            cell.alignment = Alignment(
                horizontal="right" if isinstance(display_val, float) else "left",
                vertical="center"
            )
        ws.row_dimensions[r_idx].height = 22

    # ── Xulosa qatori ──
    last_data_row = len(rows) + 2
    summary_row = last_data_row + 2
    ws.merge_cells(f"A{summary_row}:{end_col}{summary_row}")
    ws[f"A{summary_row}"] = f"📌 Xulosa: {summary}"
    ws[f"A{summary_row}"].font = Font(name="Arial", italic=True, size=10, color="444444")
    ws[f"A{summary_row}"].fill = PatternFill("solid", fgColor="FFF8E1")
    ws[f"A{summary_row}"].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    ws.row_dimensions[summary_row].height = 30

    # ── Footer ──
    footer_row = summary_row + 2
    ws.merge_cells(f"A{footer_row}:{end_col}{footer_row}")
    ws[f"A{footer_row}"] = "@slidego | t.me/slidego"
    ws[f"A{footer_row}"].font = Font(name="Arial", italic=True, size=8, color="AAAAAA")
    ws[f"A{footer_row}"].alignment = Alignment(horizontal="center")

    # ── Diagramma ──
    try:
        # Raqamli ustunlarni topish
        numeric_cols = []
        for c_idx in range(2, len(headers)+1):
            try:
                float(str(rows[0][c_idx-1]).replace(",", ".").replace(" ", ""))
                numeric_cols.append(c_idx)
            except (ValueError, IndexError, TypeError):
                pass

        if numeric_cols and len(rows) >= 3:
            chart = BarChart()
            chart.type = "col"
            chart.grouping = "clustered"
            chart.title = title
            chart.style = 10
            chart.y_axis.title = headers[numeric_cols[0]-1] if numeric_cols else ""
            chart.x_axis.title = headers[0]
            chart.width = 20; chart.height = 14

            # Birinchi raqamli ustun uchun
            data_ref = Reference(ws, min_col=numeric_cols[0], min_row=2,
                                 max_col=numeric_cols[0], max_row=last_data_row)
            cats = Reference(ws, min_col=1, min_row=3, max_row=last_data_row)
            chart.add_data(data_ref, titles_from_data=True)
            chart.set_categories(cats)

            # Agar 2 ta raqamli ustun bo'lsa, ikkinchisini ham qo'shish
            if len(numeric_cols) >= 2:
                data_ref2 = Reference(ws, min_col=numeric_cols[1], min_row=2,
                                      max_col=numeric_cols[1], max_row=last_data_row)
                chart.add_data(data_ref2, titles_from_data=True)

            ws.add_chart(chart, f"A{footer_row + 2}")
    except Exception as e:
        logger.warning(f"Diagramma yaratishda xatolik: {e}")

    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return buf


async def generate_table(topic: str, lang: str) -> BytesIO:
    data = await asyncio.to_thread(_generate_table_content, topic, lang)
    return await asyncio.to_thread(_build_excel, data, topic)


# ═══════════════════════════════════════════════
# 4. KONTSEPT XARITA (Mind Map) — PNG
# ═══════════════════════════════════════════════
def _generate_mindmap_content(topic: str, lang: str) -> dict:
    """Mind map tuzilmasini GPT dan oladi."""
    lang_inst = LANG_PROMPTS.get(lang, LANG_PROMPTS["uz"])
    client = get_client()
    prompt = (
        f"{lang_inst} {topic} mavzusida mind map tuzilmasi yarat. "
        f"JSON: {{\"center\": \"qisqa nom (max 3 so'z)\", "
        f"\"branches\": [{{\"title\": \"branch nomi (max 2 so'z)\", "
        f"\"nodes\": [\"node1 (max 3 so'z)\", \"node2\", \"node3\"]}}]}} "
        f"6 ta branch, har birida 4 ta node. Juda qisqa so'zlar."
    )
    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=900,
        temperature=0.6,
        response_format={"type": "json_object"},
    )
    return json.loads(resp.choices[0].message.content)


def _build_mindmap_png(data: dict, topic: str) -> BytesIO:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np

    center = data.get("center", topic)
    branches = data.get("branches", [])
    n = len(branches)

    # Kanvas
    fig, ax = plt.subplots(1, 1, figsize=(18, 14))
    ax.set_xlim(-11, 11); ax.set_ylim(-9, 9)
    ax.axis("off")
    fig.patch.set_facecolor("#F5F7FF")
    ax.set_facecolor("#F5F7FF")

    # Ranglar — professional
    branch_colors = [
        "#1A73E8",  # Ko'k
        "#E8711A",  # To'q sariq
        "#0F9D58",  # Yashil
        "#D93025",  # Qizil
        "#7B1FA2",  # Binafsha
        "#F4B400",  # Sariq
        "#00ACC1",  # Moviy
    ]

    # ── Markaz ──
    center_w = max(len(center) * 0.18, 3.0)
    center_h = 1.0
    ax.add_patch(mpatches.FancyBboxPatch(
        (-center_w/2, -center_h/2), center_w, center_h,
        boxstyle="round,pad=0.2", linewidth=3,
        edgecolor="#0D1B6E", facecolor="#1A1A6E", zorder=5
    ))
    ax.text(0, 0, center, ha="center", va="center",
            fontsize=14, fontweight="bold", color="white", zorder=6)

    if n == 0:
        buf = BytesIO()
        plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(); buf.seek(0)
        return buf

    # Branch joylashuvi — doira bo'ylab
    angles = np.linspace(0, 2 * np.pi, n, endpoint=False)
    # Birinchi branch yuqoridan boshlansin
    angles = angles + np.pi / 2

    # Branch masofasi
    branch_r = 5.0

    for i, (branch, angle) in enumerate(zip(branches, angles)):
        color = branch_colors[i % len(branch_colors)]
        bx = branch_r * np.cos(angle)
        by = branch_r * np.sin(angle)

        # Markaz → branch chiziq (qalin)
        ax.plot([0, bx * 0.55], [0, by * 0.55], color=color,
                linewidth=3, alpha=0.8, zorder=1, solid_capstyle='round')
        ax.plot([bx * 0.55, bx], [by * 0.55, by], color=color,
                linewidth=2, alpha=0.6, zorder=1, solid_capstyle='round')

        # Branch box
        title = branch.get("title", "")
        tw = max(len(title) * 0.16, 2.2)
        th = 0.75
        ax.add_patch(mpatches.FancyBboxPatch(
            (bx - tw/2, by - th/2), tw, th,
            boxstyle="round,pad=0.12", linewidth=2,
            edgecolor=color, facecolor=color, alpha=0.95, zorder=4
        ))
        ax.text(bx, by, title, ha="center", va="center",
                fontsize=10, fontweight="bold", color="white", zorder=5)

        # Nodes
        nodes = branch.get("nodes", [])
        nn = len(nodes)
        if nn == 0:
            continue

        # Node joylashuvi — branch atrofida yoyilgan
        # Har bir node uchun burchak hisoblash
        # Branchdan tashqariga qarab, lekin bir-biriga tegmasin
        node_r = 3.2  # Branch dan node gacha masofa
        spread = min(0.55, 0.9 / max(nn, 1))
        node_angles = np.linspace(angle - spread * (nn-1)/2,
                                  angle + spread * (nn-1)/2, nn)

        for j, (node, nangle) in enumerate(zip(nodes, node_angles)):
            nx = bx + node_r * np.cos(nangle)
            ny = by + node_r * np.sin(nangle)

            # Chegara tekshiruvi
            nx = np.clip(nx, -10.2, 10.2)
            ny = np.clip(ny, -8.2, 8.2)

            # Branch → node chiziq
            ax.plot([bx, nx], [by, ny], color=color, linewidth=1.2,
                    alpha=0.5, linestyle="-", zorder=2, solid_capstyle='round')

            # Node box
            nw = max(len(node) * 0.13, 1.8)
            nh = 0.55
            ax.add_patch(mpatches.FancyBboxPatch(
                (nx - nw/2, ny - nh/2), nw, nh,
                boxstyle="round,pad=0.08", linewidth=1.2,
                edgecolor=color, facecolor="white", alpha=0.97, zorder=3
            ))
            ax.text(nx, ny, node, ha="center", va="center",
                    fontsize=8, color="#222222", zorder=4,
                    fontfamily="DejaVu Sans")

    # ── Sarlavha ──
    ax.text(0, 8.5, f"Kontsept xarita: {topic}", ha="center", va="top",
            fontsize=15, fontweight="bold", color="#1A1A6E",
            bbox=dict(boxstyle="round,pad=0.3", facecolor="white",
                      edgecolor="#1A73E8", linewidth=1.5, alpha=0.9))

    # ── Footer ──
    ax.text(0, -8.7, "@slidego  |  t.me/slidego", ha="center", va="bottom",
            fontsize=9, color="#AAAAAA", style="italic")

    plt.tight_layout(pad=1.5)
    buf = BytesIO()
    plt.savefig(buf, format="png", dpi=160, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    plt.close()
    buf.seek(0)
    return buf


async def generate_mindmap(topic: str, lang: str) -> BytesIO:
    data = await asyncio.to_thread(_generate_mindmap_content, topic, lang)
    return await asyncio.to_thread(_build_mindmap_png, data, topic)
