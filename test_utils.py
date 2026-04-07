"""
test_utils.py — Test tuzish generatori
Berilgan mavzu bo'yicha A/B/C/D formatida:
  1. Savol varaqasi (javobsiz) — imtihon uchun
  2. Javoblar varaqasi (to'g'ri javoblar belgilangan) — o'qituvchi uchun
"""

import os
import json
import logging
from io import BytesIO
from openai import OpenAI

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


def get_client():
    return OpenAI()


# ─────────────────────────────────────────────
# Til sozlamalari
# ─────────────────────────────────────────────

LANG_PROMPTS = {
    "uz": "O'zbek tilida yoz.",
    "ru": "Пиши на русском языке.",
    "en": "Write in English.",
    "ko": "한국어로 작성하세요.",
    "zh": "用中文写。",
    "de": "Schreibe auf Deutsch.",
}

LANG_LABELS = {
    "uz": "O'zbek", "ru": "Rus", "en": "Ingliz",
    "ko": "Kores", "zh": "Xitoy", "de": "Nemis",
}

# Narx jadvali
TEST_PRICES = {
    10: 1000,
    20: 2000,
    30: 2000,
    50: 3000,
}


# ─────────────────────────────────────────────
# GPT Mega-so'rov
# ─────────────────────────────────────────────

def generate_test_content(topic: str, count: int, lang: str) -> dict:
    """Test savollarini GPT dan JSON formatida oladi."""
    lang_instruction = LANG_PROMPTS.get(lang, LANG_PROMPTS["uz"])

    prompt = f"""Sen akademik test tuzuvchisisiz. {lang_instruction}

Mavzu: "{topic}"
Savol soni: aniq {count} ta

Vazifa: Ushbu mavzu bo'yicha A/B/C/D variantli test savollarini tuz.

Quyidagi qat'iy JSON formatida qaytar:
{{
  "title": "Test sarlavhasi (masalan: Biologiya — Hujayra tuzilishi bo'yicha test)",
  "questions": [
    {{
      "number": 1,
      "question": "Savol matni?",
      "a": "A variant",
      "b": "B variant",
      "c": "C variant",
      "d": "D variant",
      "answer": "a"
    }},
    ... (jami {count} ta shunday ob'ekt bo'lishi SHART)
  ]
}}

MUHIM TALABLAR:
- Savollar soni albatta {count} ta bo'lsin.
- Har bir savolda faqat 1 ta to'g'ri javob bo'lsin.
- "answer" maydoni faqat "a", "b", "c" yoki "d" bo'lsin (kichik harf).
- Savollar mavzuning turli jihatlarini qamrab olsin.
- Variantlar bir-biridan aniq farq qilsin.
- Faqat JSON formatida qaytar."""

    client = get_client()
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {
                "role": "system",
                "content": f"Sen professional test tuzuvchisisiz. Sening vazifang {count} ta savol yaratish. Faqat JSON formatida javob ber."
            },
            {"role": "user", "content": prompt}
        ],
        response_format={"type": "json_object"},
        temperature=0.7,
        max_tokens=8000,
    )

    return json.loads(response.choices[0].message.content)


# ─────────────────────────────────────────────
# Yordamchi funksiyalar
# ─────────────────────────────────────────────

def add_divider(doc: Document, color_hex: str = "1B3A6B"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


# ─────────────────────────────────────────────
# DOCX — Savol varaqasi (javobsiz)
# ─────────────────────────────────────────────

def build_question_docx(
    content: dict,
    topic: str,
    lang: str,
    author: str = "",
    institution: str = "",
) -> BytesIO:
    """Savol varaqasi (javobsiz) DOCX yaratadi."""

    doc = Document()

    # Sahifa sozlamalari
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # ─── Sarlavha ───
    test_labels = {
        "uz": "SAVOL VARAQASI",
        "ru": "ЛИСТ ВОПРОСОВ",
        "en": "QUESTION SHEET",
        "ko": "문제지",
        "zh": "试题卷",
        "de": "FRAGEBOGEN"
    }

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(test_labels.get(lang, "SAVOL VARAQASI"))
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(27, 58, 107)

    title = content.get("title", f"{topic} bo'yicha test")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(27, 58, 107)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    add_divider(doc)

    # Muallif / muassasa
    if author or institution:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if author:
            run = p.add_run(f"F.I.O.: {author}   ")
            run.font.size = Pt(11)
        if institution:
            run = p.add_run(f"Muassasa: {institution}")
            run.font.size = Pt(11)

    # Sana va ball joyi
    date_labels = {
        "uz": "Sana: _______________   Ball: _______",
        "ru": "Дата: _______________   Балл: _______",
        "en": "Date: _______________   Score: _______",
        "ko": "날짜: _______________   점수: _______",
        "zh": "日期: _______________   分数: _______",
        "de": "Datum: _______________   Punkte: _______",
    }
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(date_labels.get(lang, "Sana: _______________   Ball: _______"))
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(12)

    add_divider(doc)
    doc.add_paragraph()

    # ─── Savollar ───
    questions = content.get("questions", [])
    option_labels = {"a": "A", "b": "B", "c": "C", "d": "D"}

    for q in questions:
        num = q.get("number", "")
        question = q.get("question", "")

        # Savol
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{num}. {question}")
        run.bold = True
        run.font.size = Pt(12)

        # Variantlar (2 ustunli)
        for opt in ["a", "b", "c", "d"]:
            opt_text = q.get(opt, "")
            if opt_text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f"{option_labels[opt]}) {opt_text}")
                run.font.size = Pt(11)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # BytesIO ga saqlash
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# DOCX — Javoblar varaqasi
# ─────────────────────────────────────────────

def build_answer_docx(
    content: dict,
    topic: str,
    lang: str,
) -> BytesIO:
    """Javoblar varaqasi (to'g'ri javoblar belgilangan) DOCX yaratadi."""

    doc = Document()

    # Sahifa sozlamalari
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # ─── Sarlavha ───
    answer_labels = {
        "uz": "JAVOBLAR VARAQASI (O'QITUVCHI UCHUN)",
        "ru": "ЛИСТ ОТВЕТОВ (ДЛЯ ПРЕПОДАВАТЕЛЯ)",
        "en": "ANSWER KEY (FOR TEACHER)",
        "ko": "정답지 (교사용)",
        "zh": "答案卷 (教师用)",
        "de": "ANTWORTBOGEN (FÜR LEHRER)"
    }

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(answer_labels.get(lang, "JAVOBLAR VARAQASI"))
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 0, 0)

    title = content.get("title", f"{topic} bo'yicha test")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(27, 58, 107)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    add_divider(doc)
    doc.add_paragraph()

    # ─── Javoblar jadvali ───
    questions = content.get("questions", [])
    option_labels = {"a": "A", "b": "B", "c": "C", "d": "D"}

    # Jadval: 5 ustun (10 ta savol bir qatorda)
    cols_per_row = 10
    rows_needed = (len(questions) + cols_per_row - 1) // cols_per_row

    # Qisqa jadval (savol raqami + javob)
    answer_header = {
        "uz": "Javoblar jadvali",
        "ru": "Таблица ответов",
        "en": "Answer Table",
        "ko": "정답표",
        "zh": "答案表",
        "de": "Antworttabelle"
    }
    p = doc.add_paragraph()
    run = p.add_run(answer_header.get(lang, "Javoblar jadvali"))
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(27, 58, 107)
    p.paragraph_format.space_after = Pt(6)

    # Jadval yaratish (2 ustun: №, Javob)
    table = doc.add_table(rows=1, cols=cols_per_row * 2)
    table.style = 'Table Grid'

    # Sarlavha qatori
    hdr_row = table.rows[0]
    for i in range(cols_per_row):
        # Savol raqami ustuni
        cell = hdr_row.cells[i * 2]
        cell.text = "№"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(cell, "1B3A6B")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # Javob ustuni
        cell2 = hdr_row.cells[i * 2 + 1]
        cell2.text = "J"
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2.paragraphs[0].runs[0].bold = True
        cell2.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(cell2, "1B3A6B")
        cell2.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Ma'lumot qatorlari
    for row_i in range(rows_needed):
        data_row = table.add_row()
        for col_i in range(cols_per_row):
            q_idx = row_i * cols_per_row + col_i
            if q_idx < len(questions):
                q = questions[q_idx]
                num = q.get("number", q_idx + 1)
                answer = q.get("answer", "").lower()
                answer_label = option_labels.get(answer, answer.upper())

                # Raqam
                cell = data_row.cells[col_i * 2]
                cell.text = str(num)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(10)

                # Javob (yashil fon)
                cell2 = data_row.cells[col_i * 2 + 1]
                cell2.text = answer_label
                cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell2.paragraphs[0].runs[0].bold = True
                cell2.paragraphs[0].runs[0].font.size = Pt(10)
                cell2.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)

    doc.add_paragraph()
    add_divider(doc)
    doc.add_paragraph()

    # ─── To'liq savollar va javoblar ───
    full_label = {
        "uz": "TO'LIQ SAVOLLAR VA JAVOBLAR",
        "ru": "ПОЛНЫЕ ВОПРОСЫ И ОТВЕТЫ",
        "en": "FULL QUESTIONS AND ANSWERS",
        "ko": "전체 문제 및 정답",
        "zh": "完整题目及答案",
        "de": "VOLLSTÄNDIGE FRAGEN UND ANTWORTEN"
    }
    p = doc.add_paragraph()
    run = p.add_run(full_label.get(lang, "TO'LIQ SAVOLLAR VA JAVOBLAR"))
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(27, 58, 107)
    p.paragraph_format.space_after = Pt(8)

    for q in questions:
        num = q.get("number", "")
        question = q.get("question", "")
        answer = q.get("answer", "").lower()

        # Savol
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{num}. {question}")
        run.bold = True
        run.font.size = Pt(11)

        # Variantlar (to'g'ri javob yashil)
        for opt in ["a", "b", "c", "d"]:
            opt_text = q.get(opt, "")
            if opt_text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_after = Pt(2)

                is_correct = (opt == answer)
                prefix = "✓ " if is_correct else "   "
                run = p.add_run(f"{prefix}{option_labels[opt]}) {opt_text}")
                run.font.size = Pt(11)
                if is_correct:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 128, 0)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # BytesIO ga saqlash
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# Asosiy funksiya
# ─────────────────────────────────────────────

async def generate_test(
    topic: str,
    count: int,
    lang: str,
    author: str = "",
    institution: str = "",
) -> tuple[BytesIO, BytesIO]:
    """
    Test yaratadi.
    Returns: (savol_varaqasi, javoblar_varaqasi) — 2 ta BytesIO
    """
    import asyncio

    # 1. GPT dan kontent olish
    content = await asyncio.to_thread(
        generate_test_content, topic, count, lang
    )

    # 2. Savol varaqasi
    question_doc = await asyncio.to_thread(
        build_question_docx, content, topic, lang, author, institution
    )

    # 3. Javoblar varaqasi
    answer_doc = await asyncio.to_thread(
        build_answer_docx, content, topic, lang
    )

    return question_doc, answer_doc
