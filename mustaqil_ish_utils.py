
import logging
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

client = OpenAI()

# --- Helper Functions ---

def add_page_border(document):
    """Hujjatning har bir sahifasiga to'rtburchak ramka qo'shadi."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for section in document.sections:
        sectPr = section._sectPr

        # Eski pgBorders ni o'chirish (agar mavjud bo'lsa)
        for old in sectPr.findall(qn('w:pgBorders')):
            sectPr.remove(old)

        pgBorders = OxmlElement('w:pgBorders')
        # offsetFrom="page" — ramka sahifa chekkasidan o'lchanadi
        pgBorders.set(qn('w:offsetFrom'), 'page')

        for edge in ('top', 'left', 'bottom', 'right'):
            border_el = OxmlElement(f'w:{edge}')
            border_el.set(qn('w:val'),   'single')
            border_el.set(qn('w:sz'),    '18')    # qalinlik (1/8 pt birligida)
            border_el.set(qn('w:space'), '24')    # sahiya chetidan masofa
            border_el.set(qn('w:color'), '000000')
            pgBorders.append(border_el)

        sectPr.append(pgBorders)


def strip_markdown(text: str) -> str:
    """GPT javobidagi markdown belgilarini tozalaydi."""
    # Sarlavha belgilari: ###, ##, #
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    # Qalin va kursiv: **text**, __text__, *text*, _text_
    text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
    text = re.sub(r'_{1,3}(.*?)_{1,3}', r'\1', text)
    # Inline kod: `text`
    text = re.sub(r'`(.*?)`', r'\1', text)
    # Gorizontal chiziq: --- yoki ***
    text = re.sub(r'^[-\*]{3,}\s*$', '', text, flags=re.MULTILINE)
    # Ortiqcha bo'sh qatorlarni birlashtirish
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def set_paragraph_font(paragraph, font_name='Times New Roman', font_size=14, bold=False, italic=False):
    """Sets font for all runs in a paragraph."""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic

def add_formatted_paragraph(document, text, font_size=14, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6), first_line_indent=None):
    """Adds a formatted paragraph to the document. If text contains \\n\\n, splits into multiple paragraphs."""
    from docx.shared import Cm
    # Avzaslarni bo'lish - \n\n bo'yicha
    parts = [p.strip() for p in text.split('\n\n') if p.strip()]
    if not parts:
        parts = [text]
    last_p = None
    for part in parts:
        p = document.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.space_after = space_after
        # Matn paragrafi uchun xat boshi chekinishi (bold bo'lmagan, justify bo'lgan)
        if first_line_indent is not None:
            p.paragraph_format.first_line_indent = first_line_indent
        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY and not bold:
            p.paragraph_format.first_line_indent = Cm(1.25)
        runner = p.add_run(part)
        runner.font.name = 'Times New Roman'
        runner.font.size = Pt(font_size)
        runner.font.bold = bold
        runner.font.italic = italic
        last_p = p
    return last_p

def generate_content_from_gpt(prompt, language, system_message):
    """Generates content using GPT and returns it as a string."""
    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": prompt}
            ],
        )
        content = response.choices[0].message.content
        return strip_markdown(content.strip())
    except Exception as e:
        logging.error(f"GPT content generation failed: {e}")
        return f"Kontent yaratishda xatolik: {e}"

# --- Document Part Creation ---

def create_cover_page(document, university_info, topic, name_surname, teacher_name, doc_type="MUSTAQIL ISH"):
    """Creates the title page for the document."""
    # ── Vazirlik (doim chiqadi, markazlashgan, qalin) ──
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("O'ZBEKISTON RESPUBLIKASI")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True

    # ── Ta'lim muassasasi (agar kiritilsa, vazirlik tagidan) ──
    if university_info:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(university_info.upper())
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.bold = True

    # ── MUSTAQIL ISH (katta, markazlashgan) ──
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(doc_type.upper())
    r.font.name = 'Times New Roman'
    r.font.size = Pt(28)
    r.bold = True

    # ── Mavzu (markazlashgan, qalin) ──
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Mavzu:")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(topic)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True

    # ── Tayyorladi (chapga, qalin, bir qatorda) ──
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Tayyorladi: ")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True
    r2 = p.add_run(name_surname or '')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(14)
    r2.bold = True

    # ── Qabul qildi (chapga, qalin, bir qatorda) ──
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Qabul qildi: ")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True
    r2 = p.add_run(teacher_name or '')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(14)
    r2.bold = True

    # ── Shahar va yil (markazlashgan, pastda) ──
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("2026")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    document.add_page_break()

def create_plan_page(document, topic, language):
    """Creates the plan page (Reja) and returns the plan items."""
    import re
    add_formatted_paragraph(document, "Reja", font_size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))

    system_msg = (
        f"You are an academic assistant. Create a plan for a research paper in {language}. "
        f"Rules: 'Kirish', 'Xulosa', 'Foydalanilgan adabiyotlar' must appear WITHOUT any number or prefix. "
        f"Main content sections (3-4 items) must be numbered starting from 1. "
        f"Return only the plain list, one item per line, no extra text."
    )
    prompt = (
        f"Mavzu: '{topic}'. Mustaqil ish rejasini tuz. "
        f"Format: birinchi qatorda 'Kirish' (raqamsiz), keyin 1., 2., 3. (yoki 4.) ta asosiy bo'lim, "
        f"so'ng 'Xulosa' (raqamsiz), oxirida 'Foydalanilgan adabiyotlar' (raqamsiz). "
        f"Faqat ro'yxatni qaytar."
    )
    plan_content = generate_content_from_gpt(prompt, language, system_msg)

    # Raqamsiz bo'lishi kerak bo'lgan kalit so'zlar
    NO_NUMBER_KEYWORDS = ["kirish", "xulosa", "foydalanilgan", "adabiyot", "introduction",
                          "conclusion", "references", "bibliography", "введение",
                          "заключение", "литература"]

    plan_items = []
    for line in plan_content.split('\n'):
        line = line.strip()
        if not line:
            continue

        # Qatorni hujjatga yozish
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(6)

        line_lower = line.lower()
        is_no_number = any(kw in line_lower for kw in NO_NUMBER_KEYWORDS)

        if is_no_number:
            # Raqamni olib tashla (agar GPT baribir qo'shgan bo'lsa)
            clean_line = re.sub(r'^[\d]+[\d\.]*\.?\s*', '', line).strip()
            r = p.add_run(clean_line)
        else:
            r = p.add_run(line)

        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)

        plan_items.append(line)

    document.add_page_break()
    return plan_items

def create_text_section(document, title, topic, prompt_detail, language):
    """Creates a standard text section like Introduction or Conclusion."""
    add_formatted_paragraph(document, title, font_size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))
    system_msg = f"You are an academic assistant writing a research paper in {language}. Write a detailed, academic paragraph."
    prompt = f"Mavzu: '{topic}'. {prompt_detail}"
    content = generate_content_from_gpt(prompt, language, system_msg)
    add_formatted_paragraph(document, content)
    document.add_page_break()

def create_main_content(document, topic, plan_items, page_count, language):
    """Creates the main content based on the plan."""
    main_plan_items = [item for item in plan_items if "kirish" not in item.lower() and "xulosa" not in item.lower() and "adabiyotlar" not in item.lower()]
    if not main_plan_items:
        return

    # Estimate words needed
    words_per_page = 300
    # Reserve pages for cover, plan, intro, conclusion, references
    main_content_pages = max(1, page_count - 5)
    words_per_item = (main_content_pages * words_per_page) // len(main_plan_items)

    system_msg = f"You are an academic assistant writing a research paper in {language}. Write a detailed, academic text of about {words_per_item} words."
    for item in main_plan_items:
        add_formatted_paragraph(document, item, font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
        prompt = f"Mavzu: '{topic}'. Rejaning quyidagi bandi bo'yicha {words_per_item} so'z atrofida batafsil ilmiy matn yozib ber: \n{item}\n\nMuhim: Matn oxirida 'Xulosa', 'Hulosa', 'Conclusion' kabi bo'lim qo'shma. Faqat shu band bo'yicha asosiy matn yoz. Har bir yangi fikrni yangi avzasdan boshlash uchun \\n\\n ishlatib avzaslarni ajrat."
        item_content = generate_content_from_gpt(prompt, language, system_msg)
        add_formatted_paragraph(document, item_content)
        document.add_paragraph() # Add space

    document.add_page_break()

def create_references_page(document, topic, language):
    """Creates the references page."""
    add_formatted_paragraph(document, "Foydalanilgan adabiyotlar", font_size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))
    system_msg = f"You are an academic assistant. Create a numbered list of 8-10 academic references in {language}, prioritizing sources from Uzbekistan."
    prompt = f"Mavzu: '{topic}'. Shu mavzu bo'yicha 8-10 ta ilmiy adabiyotlar ro'yxatini (kitoblar, maqolalar, web-saytlar) tuzib ber. O'zbekiston manbalariga ustunlik berilsin."
    content = generate_content_from_gpt(prompt, language, system_msg)
    add_formatted_paragraph(document, content, alignment=WD_ALIGN_PARAGRAPH.LEFT)

# --- Main Orchestrator Function ---

def generate_mustaqil_ish(topic, page_count, language, name_surname, university_info, teacher_name, doc_type="MUSTAQIL ISH"):
    """Generates the full 'Mustaqil ish' or 'Referat' Word document."""
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    # Har sahifaga ramka qo'shish
    add_page_border(document)

    # 1. Cover Page
    create_cover_page(document, university_info, topic, name_surname, teacher_name, doc_type=doc_type)

    # 2. Plan Page
    plan_items = create_plan_page(document, topic, language)

    # 3. Introduction Page
    create_text_section(document, "Kirish", topic, "Shu mavzu uchun mustaqil ishga 200-250 so'zdan iborat Kirish qismi yozib ber. Kirishda mavzuning dolzarbligi, maqsadi va vazifalari yoritilsin.", language)

    # 4. Main Content
    create_main_content(document, topic, plan_items, page_count, language)

    # 5. Conclusion Page
    create_text_section(document, "Xulosa", topic, "Shu mavzu bo'yicha yozilgan mustaqil ish uchun 200-250 so'zdan iborat Xulosa qismi yozib ber.", language)

    # 6. References Page
    create_references_page(document, topic, language)

    # Save to a byte stream
    doc_bytes = BytesIO()
    document.save(doc_bytes)
    doc_bytes.seek(0)

    logging.info(f"'Mustaqil ish' generated for topic: {topic}")
    return doc_bytes
