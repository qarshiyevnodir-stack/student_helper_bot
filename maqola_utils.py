"""
maqola_utils.py — Sifatli akademik/ilmiy maqola yaratuvchi modul.

Yangi tuzilma (har bo'lim alohida sahifa):
  1-sahifa  : Sarlavha + Annotatsiya (3 tilda) + Kalit so'zlar
  2-sahifa  : Kirish (alohida sahifa)
  3..N sahifa: Asosiy bo'limlar (har biri alohida sahifa) — jadval, grafik, iqtibos, rasm bilan
  N+1-sahifa: Xulosa va tavsiyalar (alohida sahifa)
  N+2-sahifa: Foydalanilgan adabiyotlar (alohida sahifa)

Sahifa tanlash: 5 / 7 / 9 / 11 / 13 / 15

OPTIMIZATSIYA: Barcha kontent BITTA mega-so'rovda (JSON) olinadi.
"""

import os
import json
import logging
import re
import tempfile
from io import BytesIO

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import requests

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
client = OpenAI()

# Ranglar
DARK_BLUE  = (31, 73, 125)
GOLD_COLOR = (180, 140, 0)
GRAY_COLOR = (89, 89, 89)

# ─────────────────────────────────────────────
# Sahifa → konfiguratsiya
# ─────────────────────────────────────────────
PAGE_CONFIG = {
    5:  {"section_count": 3,  "words_per_section": 400,  "intro_words": 350,  "conclusion_words": 300,  "ref_count": 8},
    7:  {"section_count": 5,  "words_per_section": 400,  "intro_words": 350,  "conclusion_words": 300,  "ref_count": 10},
    9:  {"section_count": 5,  "words_per_section": 500,  "intro_words": 400,  "conclusion_words": 350,  "ref_count": 12},
    11: {"section_count": 6,  "words_per_section": 550,  "intro_words": 450,  "conclusion_words": 400,  "ref_count": 14},
    13: {"section_count": 7,  "words_per_section": 580,  "intro_words": 500,  "conclusion_words": 420,  "ref_count": 16},
    15: {"section_count": 8,  "words_per_section": 600,  "intro_words": 550,  "conclusion_words": 450,  "ref_count": 18},
}


# ─────────────────────────────────────────────
# Yordamchi funksiyalar
# ─────────────────────────────────────────────

def strip_markdown(text: str) -> str:
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
    text = re.sub(r'_{1,3}(.*?)_{1,3}', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'^[-\*]{3,}\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def set_font(run, size=12, bold=False, italic=False,
             name='Times New Roman', color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc_or_cell, text='', alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  size=12, bold=False, italic=False,
                  space_before=0, space_after=6, color=None,
                  line_spacing=None):
    p = doc_or_cell.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if line_spacing:
        p.paragraph_format.line_spacing = Pt(line_spacing)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_horizontal_line(doc, color_hex='1F497D'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_page_border(document):
    for section in document.sections:
        sectPr = section._sectPr
        for old in sectPr.findall(qn('w:pgBorders')):
            sectPr.remove(old)
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for edge in ('top', 'left', 'bottom', 'right'):
            border_el = OxmlElement(f'w:{edge}')
            border_el.set(qn('w:val'),   'single')
            border_el.set(qn('w:sz'),    '12')
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), '1F497D')
            pgBorders.append(border_el)
        sectPr.append(pgBorders)


def add_section_heading(doc, number, title):
    add_paragraph(doc, f"{number}. {title.upper()}",
                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  size=15, bold=True, space_before=6, space_after=4,
                  color=DARK_BLUE)
    add_horizontal_line(doc)


def add_body_text(doc, text):
    for para in text.split('\n'):
        para = strip_markdown(para).strip()
        if para:
            add_paragraph(doc, para,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          size=14, space_before=0, space_after=6,
                          line_spacing=18)


def add_blockquote(doc, text, caption=''):
    """Iqtibos bloki — chap ko'k chiziq bilan."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.2)
    # Chap chiziq
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '1F497D')
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(f'"{strip_markdown(text)}"')
    set_font(r, size=13, italic=True, color=GRAY_COLOR)
    if caption:
        add_paragraph(doc, f"— {strip_markdown(caption)}",
                      alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                      size=11, italic=True, space_before=0, space_after=8,
                      color=GRAY_COLOR)


def add_table_to_doc(doc, table_data: dict):
    """
    table_data = {
      "caption": "Jadval 1. ...",
      "headers": ["Ustun 1", "Ustun 2", ...],
      "rows": [["qiymat", "qiymat", ...], ...]
    }
    """
    caption = table_data.get("caption", "")
    headers = table_data.get("headers", [])
    rows    = table_data.get("rows", [])

    if not headers or not rows:
        return

    if caption:
        add_paragraph(doc, caption,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      size=12, bold=True, space_before=8, space_after=4,
                      color=DARK_BLUE)

    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'

    # Sarlavha qatori
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_font(run, size=12, bold=True, color=DARK_BLUE)
        # Fon rangi
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'DCE6F1')
        tcPr.append(shd)

    # Ma'lumot qatorlari
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_val in enumerate(row_data):
            row_cells[c_idx].text = str(cell_val)
            for para in row_cells[c_idx].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    set_font(run, size=12)

    # Jadvaldan keyin bo'sh joy
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def create_chart_image(chart_data: dict) -> str | None:
    """
    chart_data = {
      "type": "bar" | "line" | "pie",
      "title": "...",
      "labels": ["A", "B", "C"],
      "values": [10, 20, 30],
      "xlabel": "...",
      "ylabel": "..."
    }
    PNG faylni temp papkaga saqlaydi va yo'lini qaytaradi.
    """
    try:
        chart_type = chart_data.get("type", "bar")
        title      = chart_data.get("title", "")
        labels     = chart_data.get("labels", [])
        values     = chart_data.get("values", [])
        xlabel     = chart_data.get("xlabel", "")
        ylabel     = chart_data.get("ylabel", "")

        if not labels or not values:
            return None

        # Qiymatlarni float ga o'tkazish
        values = [float(v) for v in values]

        fig, ax = plt.subplots(figsize=(7, 4), dpi=120)
        fig.patch.set_facecolor('#F8F9FA')
        ax.set_facecolor('#FFFFFF')

        colors = ['#1F497D', '#2E75B6', '#4472C4', '#70AD47', '#ED7D31',
                  '#FFC000', '#FF0000', '#7030A0', '#00B0F0', '#92D050']

        if chart_type == 'pie':
            wedge_colors = colors[:len(labels)]
            ax.pie(values, labels=labels, autopct='%1.1f%%',
                   colors=wedge_colors, startangle=90,
                   textprops={'fontsize': 10})
        elif chart_type == 'line':
            ax.plot(labels, values, marker='o', color='#1F497D',
                    linewidth=2, markersize=6)
            ax.fill_between(range(len(labels)), values, alpha=0.1, color='#1F497D')
            ax.set_xticks(range(len(labels)))
            ax.set_xticklabels(labels, rotation=15, ha='right', fontsize=9)
            ax.set_ylabel(ylabel, fontsize=10)
            ax.set_xlabel(xlabel, fontsize=10)
            ax.grid(axis='y', linestyle='--', alpha=0.5)
        else:  # bar
            bar_colors = colors[:len(labels)]
            bars = ax.bar(labels, values, color=bar_colors, edgecolor='white',
                          linewidth=0.5)
            for bar, val in zip(bars, values):
                ax.text(bar.get_x() + bar.get_width() / 2,
                        bar.get_height() + max(values) * 0.01,
                        f'{val:g}', ha='center', va='bottom', fontsize=9)
            ax.set_xticks(range(len(labels)))
            ax.set_xticklabels(labels, rotation=15, ha='right', fontsize=9)
            ax.set_ylabel(ylabel, fontsize=10)
            ax.set_xlabel(xlabel, fontsize=10)
            ax.grid(axis='y', linestyle='--', alpha=0.5)

        ax.set_title(title, fontsize=12, fontweight='bold',
                     color='#1F497D', pad=10)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

        plt.tight_layout()

        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        plt.savefig(tmp.name, bbox_inches='tight', dpi=120)
        plt.close(fig)
        return tmp.name
    except Exception as e:
        logging.warning(f"Grafik yaratishda xatolik: {e}")
        plt.close('all')
        return None


def fetch_topic_image(topic: str) -> str | None:
    """Mavzuga oid rasm URL dan yuklab, temp faylga saqlaydi."""
    try:
        keyword = topic.split()[0] if topic else "science"
        url = f"https://source.unsplash.com/800x400/?{keyword},research"
        resp = requests.get(url, timeout=8)
        if resp.status_code == 200 and len(resp.content) > 5000:
            tmp = tempfile.NamedTemporaryFile(suffix='.jpg', delete=False)
            tmp.write(resp.content)
            tmp.close()
            return tmp.name
    except Exception as e:
        logging.warning(f"Rasm yuklashda xatolik: {e}")
    return None


# ─────────────────────────────────────────────
# MEGA-SO'ROV
# ─────────────────────────────────────────────

def generate_all_content(topic: str, language: str,
                          article_type: str, page_count: int) -> dict:
    lang_map = {
        'uz': "o'zbek", 'ru': "rus", 'en': "ingliz",
        'ko': "kores",  'zh': "xitoy", 'de': "nemis"
    }
    type_map = {
        'ilmiy':        "ilmiy-tadqiqot",
        'publitsistik': "publitsistik",
        'tahliliy':     "tahliliy-analitik"
    }
    lang_name = lang_map.get(language, "o'zbek")
    type_name = type_map.get(article_type, "ilmiy-tadqiqot")

    cfg = PAGE_CONFIG.get(page_count, PAGE_CONFIG[5])
    section_count    = cfg["section_count"]
    word_per_section = cfg["words_per_section"]
    intro_words      = cfg["intro_words"]
    conclusion_words = cfg["conclusion_words"]
    ref_count        = cfg["ref_count"]

    system_msg = (
        f"Siz {lang_name} tilida {type_name} maqola yozuvchi yuqori malakali mutaxassississiz. "
        f"Faqat sof matn yozing, markdown belgilari ishlatmang. "
        f"Matnlar boy, to'liq va akademik uslubda bo'lsin. "
        f"Javobingiz to'liq JSON formatida bo'lsin."
    )

    # Bo'limlar uchun template — birinchi bo'limda jadval, ikkinchisida grafik, uchinchisida iqtibos
    sections_parts = []
    for i in range(section_count):
        part = f'    {{"title": "{i+1}-bo\'lim nomi ({lang_name} tilida)", "text": "kamida {word_per_section} so\'zlik boy akademik matn"'
        if i == 0:
            part += (
                ', "table": {"caption": "Jadval 1. Mavzuga oid taqqoslama ma\'lumotlar",'
                ' "headers": ["Ko\'rsatkich", "Qiymat 1", "Qiymat 2", "Qiymat 3"],'
                ' "rows": [["Birinchi", "...", "...", "..."], ["Ikkinchi", "...", "...", "..."], ["Uchinchi", "...", "...", "..."]]}'
            )
        elif i == 1:
            part += (
                ', "chart": {"type": "bar", "title": "Mavzuga oid statistika",'
                ' "labels": ["A", "B", "C", "D"], "values": [25, 40, 30, 55],'
                ' "xlabel": "Kategoriyalar", "ylabel": "Qiymatlar"}'
            )
        elif i == 2:
            part += ', "blockquote": {"text": "Mavzuga oid muhim iqtibos yoki ta\'rif", "source": "Muallif, yil"}'
        part += '}'
        sections_parts.append(part)

    sections_template = "\n".join(sections_parts)

    prompt = f"""'{topic}' mavzusida {type_name} maqola uchun quyidagi JSON strukturasini to'ldiring.
Til: {lang_name}. Barcha matnlar {lang_name} tilida bo'lsin.
Maqola turi: {type_name}. Tanlangan hajm: {page_count} sahifa.
MATNLAR JUDA BOY VA TO'LIQ BO'LSIN.

{{
  "title": "Maqolaning rasmiy sarlavhasi ({lang_name} tilida)",
  "annotation_uz": "100-120 so'zlik annotatsiya o'zbek tilida",
  "annotation_ru": "100-120 so'zlik annotatsiya rus tilida",
  "annotation_en": "100-120 so'zlik annotatsiya ingliz tilida (Abstract)",
  "keywords": ["kalit so'z 1", "kalit so'z 2", "kalit so'z 3", "kalit so'z 4", "kalit so'z 5", "kalit so'z 6"],
  "introduction": "kamida {intro_words} so'zlik kirish — dolzarblik, maqsad, vazifalar, metodologiya",
  "sections": [
{sections_template}
  ],
  "conclusion": "kamida {conclusion_words} so'zlik xulosa — asosiy natijalar va ilmiy hissa",
  "recommendations": "kamida 150 so'zlik amaliy tavsiyalar",
  "references": [
    "1. Birinchi adabiyot (APA formatida)",
    "... ({ref_count} ta manba)"
  ]
}}

Faqat JSON qaytaring. Bo'limlar soni: {section_count} ta.
Jadval, grafik va iqtibos ma'lumotlarini mavzuga mos HAQIQIY raqamlar va ma'lumotlar bilan to'ldiring.
Grafik values faqat raqamlar (sonlar) bo'lsin."""

    try:
        resp = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user",   "content": prompt},
            ],
            response_format={"type": "json_object"},
        )
        raw = resp.choices[0].message.content.strip()
        data = json.loads(raw)
        logging.info(f"Maqola mega-so'rov OK: {len(data.get('sections', []))} bo'lim")
        return data
    except Exception as e:
        logging.error(f"Maqola mega-so'rov xatolik: {e}")
        return {
            "title": topic,
            "annotation_uz": "Annotatsiya yaratishda xatolik.",
            "annotation_ru": "Ошибка при создании аннотации.",
            "annotation_en": "Error generating annotation.",
            "keywords": [topic],
            "introduction": "Kirish yaratishda xatolik.",
            "sections": [{"title": f"Bo'lim {i+1}", "text": "Matn yaratishda xatolik."} for i in range(3)],
            "conclusion": "Xulosa yaratishda xatolik.",
            "recommendations": "Tavsiyalar yaratishda xatolik.",
            "references": ["1. -"]
        }


# ─────────────────────────────────────────────
# DOCX YARATISH
# ─────────────────────────────────────────────

def build_maqola_docx(content: dict, topic: str, language: str,
                       article_type: str, page_count: int,
                       name_surname: str = '', university: str = '') -> BytesIO:
    type_map = {
        'ilmiy':        "Ilmiy maqola",
        'publitsistik': "Publitsistik maqola",
        'tahliliy':     "Tahliliy maqola"
    }
    type_display = type_map.get(article_type, "Ilmiy maqola")

    doc = Document()

    for section in doc.sections:
        section.page_width    = Inches(8.27)
        section.page_height   = Inches(11.69)
        section.left_margin   = Inches(1.18)
        section.right_margin  = Inches(0.79)
        section.top_margin    = Inches(0.98)
        section.bottom_margin = Inches(0.98)

    add_page_border(doc)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    # ══════════════════════════════════════════
    # 1-SAHIFA: Sarlavha + Annotatsiya + Kalit so'zlar
    # ══════════════════════════════════════════
    add_paragraph(doc, type_display.upper(),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, bold=False, space_before=0, space_after=4,
                  color=DARK_BLUE)

    title = content.get("title", topic)
    add_paragraph(doc, title,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  size=18, bold=True, space_before=8, space_after=6,
                  color=DARK_BLUE)

    add_horizontal_line(doc)

    if name_surname and name_surname.strip():
        add_paragraph(doc, name_surname.strip(),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      size=13, bold=True, space_before=8, space_after=2)
    if university and university.strip():
        add_paragraph(doc, university.strip(),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      size=12, italic=True, space_before=0, space_after=2,
                      color=GRAY_COLOR)

    add_horizontal_line(doc)

    ann_uz = content.get("annotation_uz", content.get("annotation", ""))
    ann_ru = content.get("annotation_ru", "")
    ann_en = content.get("annotation_en", "")

    for ann_label, ann_text in [
        ("ANNOTATSIYA (O'zbek)", ann_uz),
        ("АННОТАЦИЯ (Русский)", ann_ru),
        ("ABSTRACT (English)", ann_en),
    ]:
        if ann_text and ann_text.strip():
            add_paragraph(doc, ann_label,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          size=11, bold=True, space_before=8, space_after=2,
                          color=DARK_BLUE)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            r = p.add_run(strip_markdown(ann_text))
            set_font(r, size=11, italic=True, color=GRAY_COLOR)

    keywords = content.get("keywords", [])
    if keywords:
        kw_text = ", ".join(str(k) for k in keywords if k)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        r1 = p.add_run("Kalit so'zlar / Keywords: ")
        set_font(r1, size=11, bold=True, color=DARK_BLUE)
        r2 = p.add_run(kw_text)
        set_font(r2, size=11, italic=True)

    # ══════════════════════════════════════════
    # 2-SAHIFA: Kirish
    # ══════════════════════════════════════════
    doc.add_page_break()
    add_section_heading(doc, 1, "KIRISH")

    # Kirish sahifasiga rasm qo'shish
    img_path = fetch_topic_image(topic)
    if img_path:
        try:
            doc.add_picture(img_path, width=Inches(5.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_paragraph(doc, f"Rasm 1. {strip_markdown(title)}",
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          size=10, italic=True, space_before=2, space_after=8,
                          color=GRAY_COLOR)
        except Exception as e:
            logging.warning(f"Rasm qo'shishda xatolik: {e}")
        finally:
            try:
                os.unlink(img_path)
            except Exception:
                pass

    introduction = content.get("introduction", "")
    add_body_text(doc, introduction)

    # ══════════════════════════════════════════
    # 3..N-SAHIFA: Asosiy bo'limlar
    # ══════════════════════════════════════════
    sections = content.get("sections", [])
    for i, sec in enumerate(sections):
        doc.add_page_break()
        sec_title = strip_markdown(str(sec.get("title", f"Bo'lim {i+1}"))).strip()
        sec_text  = strip_markdown(str(sec.get("text", ""))).strip()

        add_section_heading(doc, i + 2, sec_title)
        add_body_text(doc, sec_text)

        # Jadval
        table_data = sec.get("table")
        if table_data and isinstance(table_data, dict):
            try:
                add_table_to_doc(doc, table_data)
            except Exception as e:
                logging.warning(f"Jadval qo'shishda xatolik: {e}")

        # Grafik
        chart_data = sec.get("chart")
        if chart_data and isinstance(chart_data, dict):
            chart_path = create_chart_image(chart_data)
            if chart_path:
                try:
                    doc.add_picture(chart_path, width=Inches(5.5))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    chart_title = chart_data.get("title", "Diagramma")
                    add_paragraph(doc, f"Diagramma {i+1}. {chart_title}",
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  size=10, italic=True, space_before=2, space_after=8,
                                  color=GRAY_COLOR)
                except Exception as e:
                    logging.warning(f"Grafik qo'shishda xatolik: {e}")
                finally:
                    try:
                        os.unlink(chart_path)
                    except Exception:
                        pass

        # Iqtibos
        blockquote = sec.get("blockquote")
        if blockquote and isinstance(blockquote, dict):
            bq_text   = blockquote.get("text", "")
            bq_source = blockquote.get("source", "")
            if bq_text:
                add_blockquote(doc, bq_text, bq_source)

    # ══════════════════════════════════════════
    # N+1-SAHIFA: Xulosa va tavsiyalar
    # ══════════════════════════════════════════
    doc.add_page_break()
    xulosa_num = len(sections) + 2
    add_section_heading(doc, xulosa_num, "XULOSA VA TAVSIYALAR")

    conclusion = content.get("conclusion", "")
    add_body_text(doc, conclusion)

    recommendations = content.get("recommendations", "")
    if recommendations and recommendations.strip():
        add_paragraph(doc, "Tavsiyalar:",
                      alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      size=14, bold=True, space_before=10, space_after=4,
                      color=DARK_BLUE)
        add_body_text(doc, recommendations)

    # ══════════════════════════════════════════
    # N+2-SAHIFA: Adabiyotlar
    # ══════════════════════════════════════════
    doc.add_page_break()
    add_paragraph(doc, "FOYDALANILGAN ADABIYOTLAR",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  size=15, bold=True, space_before=6, space_after=6,
                  color=DARK_BLUE)
    add_horizontal_line(doc)

    references = content.get("references", [])
    for ref in references:
        ref_text = strip_markdown(str(ref)).strip()
        if ref_text and not ref_text.startswith("..."):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(5)
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            r = p.add_run(ref_text)
            set_font(r, size=12)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# Asosiy generator
# ─────────────────────────────────────────────

def generate_maqola(topic: str, language: str, article_type: str,
                     page_count: int, name_surname: str = '',
                     university: str = '') -> BytesIO:
    logging.info(f"Maqola: '{topic}' | {language} | {article_type} | {page_count} sah")
    content = generate_all_content(topic, language, article_type, page_count)
    return build_maqola_docx(
        content=content, topic=topic, language=language,
        article_type=article_type, page_count=page_count,
        name_surname=name_surname, university=university,
    )
