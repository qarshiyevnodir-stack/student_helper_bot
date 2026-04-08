"""
insho_utils.py — Insho / Esse generatori
Akademik va professional uslubda insho/esse yozadi.
DOCX formatida chiqaradi — Times New Roman 12pt, 1.5 interval.
"""
import os
import logging
from io import BytesIO
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


def get_client():
    return OpenAI()


# ─────────────────────────────────────────────
# Til sozlamalari
# ─────────────────────────────────────────────
LANG_PROMPTS = {
    "uz": "O'zbek tilida yoz. Akademik, rasmiy uslubda.",
    "ru": "Пиши на русском языке. Академический, официальный стиль.",
    "en": "Write in English. Use academic, formal style.",
    "ko": "한국어로 작성하세요. 학술적이고 공식적인 문체를 사용하세요.",
    "zh": "用中文写。使用学术、正式的风格。",
    "de": "Schreibe auf Deutsch. Verwende einen akademischen, formellen Stil.",
}

LANG_LABELS = {
    "uz": "O'zbek",
    "ru": "Rus",
    "en": "Ingliz",
    "ko": "Kores",
    "zh": "Xitoy",
    "de": "Nemis",
}

# Insho turlari
INSHO_TYPES = {
    "erkin": "Erkin insho",
    "tahliliy": "Tahliliy esse",
    "argumentativ": "Argumentativ esse",
    "tavsifiy": "Tavsifiy insho",
    "muqoyasali": "Muqoyasali esse",
}

INSHO_TYPE_LABELS = {
    "erkin":        "✍️ Erkin insho",
    "tahliliy":     "🔍 Tahliliy esse",
    "argumentativ": "💡 Argumentativ esse",
    "tavsifiy":     "📖 Tavsifiy insho",
    "muqoyasali":   "⚖️ Muqoyasali esse",
}

# Sahifa sozlamalari — 1 A4 sahifa ≈ 400 so'z (12pt, 1.5 interval)
INSHO_PAGES = {
    1: {"intro_words": 120, "body_paragraphs": 2, "body_words": 200, "conclusion_words": 80},
    2: {"intro_words": 180, "body_paragraphs": 3, "body_words": 320, "conclusion_words": 120},
    3: {"intro_words": 220, "body_paragraphs": 4, "body_words": 500, "conclusion_words": 150},
    5: {"intro_words": 300, "body_paragraphs": 6, "body_words": 900, "conclusion_words": 200},
}

INSHO_PRICES = {
    1: 1000,
    2: 2000,
    3: 2000,
    5: 3000,
}

# ─────────────────────────────────────────────
# GPT so'rovi — kontent generatsiya
# ─────────────────────────────────────────────
def generate_insho_content(
    topic: str,
    insho_type: str,
    lang: str,
    pages: int,
    author: str = "",
    institution: str = "",
) -> dict:
    """Insho kontentini GPT orqali yaratadi."""
    cfg = INSHO_PAGES.get(pages, INSHO_PAGES[2])
    lang_instruction = LANG_PROMPTS.get(lang, LANG_PROMPTS["uz"])
    type_label = INSHO_TYPES.get(insho_type, "Erkin insho")

    intro_words = cfg["intro_words"]
    body_paragraphs = cfg["body_paragraphs"]
    body_words = cfg["body_words"]
    conclusion_words = cfg["conclusion_words"]

    total_words = intro_words + body_words + conclusion_words

    # Insho turi bo'yicha yo'riqnoma
    type_instructions = {
        "erkin": (
            "Erkin insho: Mavzu haqida shaxsiy fikr, tajriba va kuzatishlar asosida "
            "yozilgan, lekin akademik uslubda. Mantiqiy tuzilma saqlangan bo'lsin."
        ),
        "tahliliy": (
            "Tahliliy esse: Mavzuni chuqur tahlil qil. Sabab-natija bog'liqligini "
            "ko'rsat. Dalillar va misollar keltir. Ilmiy adabiyotlarga murojaat qil."
        ),
        "argumentativ": (
            "Argumentativ esse: Aniq bir pozitsiyani himoya qil. Kamida 3 ta kuchli "
            "argument keltir. Qarshi fikrlarni ham ko'rib, ularni rad et. "
            "Xulosada o'z pozitsiyangni mustahkamla."
        ),
        "tavsifiy": (
            "Tavsifiy insho: Mavzuni batafsil tasvirla. Ko'rgazmali, aniq va "
            "ifodali til ishlat. O'quvchida aniq tasavvur hosil qil."
        ),
        "muqoyasali": (
            "Muqoyasali esse: Ikki yoki undan ortiq hodisa, tushuncha yoki "
            "yondashuvni solishtir. O'xshashlik va farqlarni ko'rsat. "
            "Xulosada qaysi biri afzalroq ekanligini asosla."
        ),
    }

    type_guide = type_instructions.get(insho_type, type_instructions["erkin"])

    system_prompt = (
        f"Sen professional akademik yozuvchisan. "
        f"{lang_instruction} "
        f"Har doim to'liq, sifatli, akademik uslubda yoz. "
        f"Hech qachon qisqartirma. Berilgan so'z sonini ALBATTA bajara."
    )

    user_prompt = f"""Quyidagi mavzuda {type_label} yoz:

MAVZU: {topic}
TUR: {type_label}
{f"MUALLIF: {author}" if author else ""}
{f"MUASSASA: {institution}" if institution else ""}

YO'RIQNOMA: {type_guide}

TUZILMA (qat'iy bajar):
1. KIRISH — KAMIDA {intro_words} so'z
   - Mavzuning dolzarbligi va ahamiyati
   - Asosiy tezis (maqsad)
   - Insho tuzilmasi haqida qisqacha

2. ASOSIY QISM — {body_paragraphs} ta paragraf, JAMI KAMIDA {body_words} so'z
   - Har paragraf: aniq fikr + dalil + misol + tahlil
   - Paragraflar mantiqiy bog'langan bo'lsin
   - Akademik atamalar va ilmiy yondashuv

3. XULOSA — KAMIDA {conclusion_words} so'z
   - Asosiy fikrlarni umumlashtirish
   - Amaliy tavsiyalar
   - Yakuniy xulosaviy fikr

JAMI: KAMIDA {total_words} so'z yoz. Bu MAJBURIY.

JSON formatida qaytargin:
{{
  "title": "Insho sarlavhasi",
  "kirish": "To'liq kirish matni ({intro_words}+ so'z)",
  "paragraphs": [
    {{"heading": "Paragraf sarlavhasi", "text": "Paragraf matni (kamida {body_words // body_paragraphs} so'z)"}},
    ...
  ],
  "xulosa": "To'liq xulosa matni ({conclusion_words}+ so'z)"
}}"""

    client = get_client()
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        max_tokens=8000,
        temperature=0.75,
        response_format={"type": "json_object"},
    )

    import json
    raw = response.choices[0].message.content
    data = json.loads(raw)
    return data


def expand_section(text: str, min_words: int, lang: str, topic: str) -> str:
    """Agar matn qisqa bo'lsa, kengaytiradi."""
    word_count = len(text.split())
    if word_count >= int(min_words * 0.85):
        return text

    lang_instruction = LANG_PROMPTS.get(lang, LANG_PROMPTS["uz"])
    client = get_client()
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    f"Sen akademik yozuvchisan. {lang_instruction} "
                    f"Berilgan matnni kengaytir, boyit, lekin ma'nosini o'zgartirma."
                ),
            },
            {
                "role": "user",
                "content": (
                    f"Quyidagi matnni KAMIDA {min_words} so'zgacha kengaytir. "
                    f"Mavzu: {topic}. Faqat kengaytirilgan matnni qaytargin:\n\n{text}"
                ),
            },
        ],
        max_tokens=3000,
        temperature=0.7,
    )
    return response.choices[0].message.content.strip()


# ─────────────────────────────────────────────
# DOCX yaratish
# ─────────────────────────────────────────────
def set_paragraph_format(para, font_size=12, bold=False, italic=False,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=6, space_after=6,
                          line_spacing=1.5, color=None):
    """Paragrafga standart formatlash qo'llaydi."""
    para.alignment = alignment
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing

    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

    # Font fallback
    for run in para.runs:
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rPr.insert(0, rFonts)


def add_heading(doc: Document, text: str, level: int = 1):
    """Sarlavha qo'shadi."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16 if level == 1 else 13)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    else:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
    para.paragraph_format.space_before = Pt(12 if level == 1 else 10)
    para.paragraph_format.space_after = Pt(8)
    return para


def add_section_title(doc: Document, text: str):
    """Bo'lim sarlavhasini qo'shadi."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text.upper())
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(6)
    # Pastki chiziq
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A73E8')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def add_body_paragraph(doc: Document, text: str, first_indent: bool = True):
    """Asosiy matn paragrafini qo'shadi."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_indent:
        para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = 1.5

    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.insert(0, rFonts)
    return para


def build_insho_docx(
    data: dict,
    topic: str,
    insho_type: str,
    lang: str,
    pages: int,
    author: str = "",
    institution: str = "",
) -> BytesIO:
    """Insho DOCX hujjatini yaratadi."""
    doc = Document()

    # Sahifa sozlamalari — A4, standart chegaralar
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── Sarlavha sahifasi ──
    title = data.get("title", topic)
    type_label = INSHO_TYPES.get(insho_type, "Insho")

    # Muassasa
    if institution:
        inst_para = doc.add_paragraph()
        inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inst_run = inst_para.add_run(institution.upper())
        inst_run.font.name = "Times New Roman"
        inst_run.font.size = Pt(11)
        inst_run.bold = True
        inst_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        inst_para.paragraph_format.space_after = Pt(4)

    # Tur belgisi
    type_para = doc.add_paragraph()
    type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_run = type_para.add_run(type_label.upper())
    type_run.font.name = "Times New Roman"
    type_run.font.size = Pt(11)
    type_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    type_para.paragraph_format.space_before = Pt(30)
    type_para.paragraph_format.space_after = Pt(8)

    # Asosiy sarlavha
    add_heading(doc, title, level=1)

    # Muallif
    if author:
        auth_para = doc.add_paragraph()
        auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        auth_run = auth_para.add_run(f"Muallif: {author}")
        auth_run.font.name = "Times New Roman"
        auth_run.font.size = Pt(12)
        auth_run.italic = True
        auth_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        auth_para.paragraph_format.space_before = Pt(6)
        auth_para.paragraph_format.space_after = Pt(4)

    # Til va sahifa
    lang_label = LANG_LABELS.get(lang, lang)
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f"Til: {lang_label}  |  Hajm: {pages} sahifa")
    info_run.font.name = "Times New Roman"
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    info_para.paragraph_format.space_after = Pt(20)

    # Chiziq
    hr_para = doc.add_paragraph()
    hr_para.paragraph_format.space_before = Pt(0)
    hr_para.paragraph_format.space_after = Pt(20)
    pPr = hr_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A73E8')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── KIRISH ──
    add_section_title(doc, "Kirish")
    kirish = data.get("kirish", "")
    for para_text in kirish.split("\n"):
        para_text = para_text.strip()
        if para_text:
            add_body_paragraph(doc, para_text)

    # ── ASOSIY QISM ──
    paragraphs = data.get("paragraphs", [])
    if paragraphs:
        add_section_title(doc, "Asosiy qism")
        for i, para_data in enumerate(paragraphs, 1):
            heading = para_data.get("heading", f"{i}-paragraf")
            text = para_data.get("text", "")

            # Paragraf sarlavhasi
            ph_para = doc.add_paragraph()
            ph_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ph_run = ph_para.add_run(f"{i}. {heading}")
            ph_run.bold = True
            ph_run.font.name = "Times New Roman"
            ph_run.font.size = Pt(12)
            ph_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
            ph_para.paragraph_format.space_before = Pt(10)
            ph_para.paragraph_format.space_after = Pt(4)

            # Paragraf matni
            for line in text.split("\n"):
                line = line.strip()
                if line:
                    add_body_paragraph(doc, line)

    # ── XULOSA ──
    add_section_title(doc, "Xulosa")
    xulosa = data.get("xulosa", "")
    for para_text in xulosa.split("\n"):
        para_text = para_text.strip()
        if para_text:
            add_body_paragraph(doc, para_text)

    # ── Footer ──
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Biz bilan ishingiz oson! | @slidego | t.me/slidego")
    footer_run.font.size = Pt(9)
    footer_run.font.name = "Times New Roman"
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# Asosiy funksiya
# ─────────────────────────────────────────────
async def generate_insho(
    topic: str,
    insho_type: str,
    lang: str,
    pages: int,
    author: str = "",
    institution: str = "",
) -> BytesIO:
    """Insho yaratadi va BytesIO qaytaradi."""
    import asyncio

    data = await asyncio.to_thread(
        generate_insho_content,
        topic, insho_type, lang, pages, author, institution
    )

    cfg = INSHO_PAGES.get(pages, INSHO_PAGES[2])

    # Kirish yetarli uzunlikdami tekshirish
    kirish = data.get("kirish", "")
    if len(kirish.split()) < int(cfg["intro_words"] * 0.8):
        kirish = await asyncio.to_thread(
            expand_section, kirish, cfg["intro_words"], lang, topic
        )
        data["kirish"] = kirish

    # Xulosa yetarli uzunlikdami tekshirish
    xulosa = data.get("xulosa", "")
    if len(xulosa.split()) < int(cfg["conclusion_words"] * 0.8):
        xulosa = await asyncio.to_thread(
            expand_section, xulosa, cfg["conclusion_words"], lang, topic
        )
        data["xulosa"] = xulosa

    doc_bytes = await asyncio.to_thread(
        build_insho_docx,
        data, topic, insho_type, lang, pages, author, institution
    )
    return doc_bytes
