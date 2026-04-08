"""
annotatsiya_utils.py — Annotatsiya xizmati
Mavzu/sarlavha asosida 100-150 so'zlik akademik annotatsiya DOCX yaratadi.
Optimizatsiyalangan: 1 ta GPT so'rov, 5-8 soniya.
"""
import asyncio
import logging
from io import BytesIO
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

def get_client():
    return OpenAI()

ANNOTATSIYA_PRICE = 1000

LANG_PROMPTS = {
    "uz": "O'zbek tilida yoz.",
    "ru": "Пиши на русском.",
    "en": "Write in English.",
    "ko": "한국어로 작성하세요.",
    "zh": "用中文写。",
    "de": "Schreibe auf Deutsch.",
}

LANG_LABELS = {
    "uz": "O'zbek", "ru": "Rus", "en": "Ingliz",
    "ko": "Kores", "zh": "Xitoy", "de": "Nemis",
}

ANNOTATSIYA_TYPES = {
    "uz": {
        "ilmiy":    "📚 Ilmiy maqola",
        "kurs":     "📄 Kurs ishi",
        "kitob":    "📖 Kitob",
        "diplom":   "🎓 Diplom ishi",
        "referat":  "📋 Referat",
    },
    "ru": {
        "ilmiy":    "📚 Научная статья",
        "kurs":     "📄 Курсовая работа",
        "kitob":    "📖 Книга",
        "diplom":   "🎓 Дипломная работа",
        "referat":  "📋 Реферат",
    },
    "en": {
        "ilmiy":    "📚 Research Article",
        "kurs":     "📄 Course Work",
        "kitob":    "📖 Book",
        "diplom":   "🎓 Thesis",
        "referat":  "📋 Essay",
    },
}

TYPE_LABELS_UZ = {
    "ilmiy":   "Ilmiy maqola",
    "kurs":    "Kurs ishi",
    "kitob":   "Kitob",
    "diplom":  "Diplom ishi",
    "referat": "Referat",
}

def _generate_annotation(title: str, doc_type: str, lang: str, author: str = "") -> str:
    lang_inst = LANG_PROMPTS.get(lang, LANG_PROMPTS["uz"])
    type_label = TYPE_LABELS_UZ.get(doc_type, doc_type)
    client = get_client()

    prompt = (
        f"{lang_inst} Akademik annotatsiya yoz. "
        f"Asar turi: {type_label}. "
        f"Sarlavha: {title}. "
        f"{('Muallif: ' + author + '.') if author else ''} "
        f"Tuzilma: 1) Asarning maqsadi va dolzarbligi, "
        f"2) Asosiy mavzu va usullar, "
        f"3) Asosiy natijalar va ahamiyati. "
        f"Hajm: 120-150 so'z. Rasmiy akademik uslub. "
        f"Faqat annotatsiya matnini yoz, sarlavha yozma."
    )
    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=600,
        temperature=0.6,
    )
    return resp.choices[0].message.content.strip()


def _build_annotation_docx(text: str, title: str, doc_type: str,
                            author: str, lang: str) -> BytesIO:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)

    def add_run(para, txt, bold=False, size=12, color=None, italic=False):
        r = para.add_run(txt)
        r.bold = bold; r.italic = italic
        r.font.name = "Times New Roman"; r.font.size = Pt(size)
        if color: r.font.color.rgb = color
        return r

    # ── Sarlavha ──
    title_map = {
        "uz": "ANNOTATSIYA", "ru": "АННОТАЦИЯ",
        "en": "ABSTRACT", "ko": "초록", "zh": "摘要", "de": "ZUSAMMENFASSUNG"
    }
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title_map.get(lang, "ANNOTATSIYA"))
    tr.bold = True; tr.font.name = "Times New Roman"
    tr.font.size = Pt(16); tr.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    tp.paragraph_format.space_after = Pt(6)

    # ── Asar nomi ──
    np_ = doc.add_paragraph()
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = np_.add_run(title)
    nr.bold = True; nr.font.name = "Times New Roman"
    nr.font.size = Pt(13); nr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
    np_.paragraph_format.space_after = Pt(4)

    # ── Muallif va tur ──
    type_label = TYPE_LABELS_UZ.get(doc_type, doc_type)
    meta_parts = []
    if author: meta_parts.append(author)
    meta_parts.append(type_label)
    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = mp.add_run("  |  ".join(meta_parts))
    mr.italic = True; mr.font.name = "Times New Roman"
    mr.font.size = Pt(11); mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    mp.paragraph_format.space_after = Pt(16)

    # ── Ajratuvchi chiziq ──
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(14)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), '1A73E8')
    pBdr.append(b); pPr.append(pBdr)

    # ── Annotatsiya matni ──
    for para_text in text.split("\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(para_text)
        r.font.name = "Times New Roman"; r.font.size = Pt(12)

    # ── Kalit so'zlar bo'limi (bo'sh, foydalanuvchi to'ldiradi) ──
    doc.add_paragraph()
    kp = doc.add_paragraph()
    kp.paragraph_format.space_before = Pt(10)
    kp.paragraph_format.space_after = Pt(4)
    kw_map = {
        "uz": "Kalit so'zlar:", "ru": "Ключевые слова:",
        "en": "Keywords:", "ko": "주요어:", "zh": "关键词:", "de": "Schlüsselwörter:"
    }
    kr = kp.add_run(kw_map.get(lang, "Kalit so'zlar:"))
    kr.bold = True; kr.font.name = "Times New Roman"; kr.font.size = Pt(12)
    kr2 = kp.add_run("  _______________________________________________")
    kr2.font.name = "Times New Roman"; kr2.font.size = Pt(12)
    kr2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Footer ──
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(30)
    fr = fp.add_run("@slidego | t.me/slidego")
    fr.font.name = "Times New Roman"; fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    fr.italic = True

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf


async def generate_annotation(title: str, doc_type: str, lang: str,
                               author: str = "") -> BytesIO:
    text = await asyncio.to_thread(_generate_annotation, title, doc_type, lang, author)
    return await asyncio.to_thread(_build_annotation_docx, text, title, doc_type, author, lang)
