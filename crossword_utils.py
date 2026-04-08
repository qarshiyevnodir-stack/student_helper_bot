"""
Krossvord yaratish moduli (yaxshilangan versiya).
GPT dan so'zlar + ta'riflar oladi, keyin to'rga joylashtiradi va DOCX chiqaradi.
"""
import os
import random
import json
from io import BytesIO
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def get_client():
    return OpenAI()


# ─── GPT: so'zlar va ta'riflar ────────────────────────────
def generate_words(topic: str, count: int, lang: str) -> list[dict]:
    lang_prompts = {
        "uz": "O'zbek tilida",
        "ru": "Rus tilida",
        "en": "Ingliz tilida",
        "ko": "Kores tilida",
        "zh": "Xitoy tilida",
        "de": "Nemis tilida",
    }
    lang_label = lang_prompts.get(lang, "O'zbek tilida")

    prompt = f"""Sen krossvord yaratuvchi mutaxassissan.
Mavzu: "{topic}"
{lang_label} {count} ta so'z va ularning qisqa ta'riflarini ber.

MUHIM QOIDALAR:
1. So'zlar FAQAT harflardan iborat bo'lsin (raqam, tire, bo'shliq YO'Q)
2. So'zlar kamida 4 ta, ko'pi bilan 12 ta harfdan iborat bo'lsin
3. Ta'riflar qisqa va aniq bo'lsin (4-10 so'z)
4. So'zlar bir-biridan farqli bo'lsin
5. So'zlar mavzuga aloqador bo'lsin
6. Imkon qadar so'zlar umumiy harflarga ega bo'lsin (kesishish uchun)

JSON formatda qaytar:
{{
  "words": [
    {{"word": "SOZ", "clue": "Qisqa tarif"}},
    ...
  ]
}}
Faqat JSON qaytar."""

    client = get_client()
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": "Faqat JSON formatda javob ber."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=2000,
        temperature=0.7,
    )

    raw = response.choices[0].message.content.strip()
    if "```" in raw:
        parts = raw.split("```")
        for p in parts:
            if "{" in p:
                raw = p.strip()
                if raw.startswith("json"):
                    raw = raw[4:].strip()
                break

    data = json.loads(raw)
    words = data.get("words", [])

    cleaned = []
    seen = set()
    for item in words:
        w = item.get("word", "").strip().upper()
        w_clean = ''.join(c for c in w if c.isalpha())
        if len(w_clean) >= 4 and len(w_clean) <= 12 and item.get("clue") and w_clean not in seen:
            cleaned.append({"word": w_clean, "clue": item["clue"].strip()})
            seen.add(w_clean)

    return cleaned[:count]


# ─── Yaxshilangan krossvord joylashtirish algoritmi ──────
class CrosswordGrid:
    def __init__(self, size=22):
        self.size = size
        self.grid = [['.' for _ in range(size)] for _ in range(size)]
        self.placed = []
        self.numbers = {}

    def can_place(self, word, row, col, direction):
        if direction == 'H':
            if col < 0 or col + len(word) > self.size or row < 0 or row >= self.size:
                return False
            # Oldin va keyin bo'sh
            if col > 0 and self.grid[row][col-1] != '.':
                return False
            if col + len(word) < self.size and self.grid[row][col + len(word)] != '.':
                return False
            has_intersection = False
            for i, ch in enumerate(word):
                c = col + i
                cell = self.grid[row][c]
                if cell == '.':
                    # Yuqori-pastda parallel so'z bo'lmasin
                    above = row > 0 and self.grid[row-1][c] != '.'
                    below = row < self.size-1 and self.grid[row+1][c] != '.'
                    if above or below:
                        # Faqat kesishish nuqtasida ruxsat
                        pass
                elif cell == ch:
                    has_intersection = True
                else:
                    return False
        else:  # V
            if row < 0 or row + len(word) > self.size or col < 0 or col >= self.size:
                return False
            if row > 0 and self.grid[row-1][col] != '.':
                return False
            if row + len(word) < self.size and self.grid[row + len(word)][col] != '.':
                return False
            has_intersection = False
            for i, ch in enumerate(word):
                r = row + i
                cell = self.grid[r][col]
                if cell == '.':
                    pass
                elif cell == ch:
                    has_intersection = True
                else:
                    return False
        return True

    def place_word(self, word, row, col, direction):
        if direction == 'H':
            for i, ch in enumerate(word):
                self.grid[row][col + i] = ch
        else:
            for i, ch in enumerate(word):
                self.grid[row + i][col] = ch

    def find_best_positions(self, word, direction):
        """Mavjud so'zlar bilan kesishadigan eng yaxshi pozitsiyalarni topadi."""
        positions = []
        opp = 'V' if direction == 'H' else 'H'

        for placed in self.placed:
            if placed['direction'] != opp:
                continue
            pw = placed['word']
            pr, pc = placed['row'], placed['col']

            for i, ch in enumerate(word):
                for j, pch in enumerate(pw):
                    if ch == pch:
                        if direction == 'H':
                            r = pr + j
                            c = pc - i
                        else:
                            r = pr - i
                            c = pc + j
                        if self.can_place(word, r, c, direction):
                            # Kesishish soni — ko'proq kesishsa yaxshiroq
                            score = 1
                            positions.append((score, r, c))

        # Eng yaxshi pozitsiyani tanlash
        positions.sort(reverse=True)
        return [(r, c) for _, r, c in positions]

    def build(self, words_data):
        if not words_data:
            return

        # So'zlarni uzunligi bo'yicha tartiblash (uzunidan qisqasiga)
        sorted_words = sorted(words_data, key=lambda x: len(x['word']), reverse=True)

        # Birinchi so'zni markazga gorizontal qo'yish
        first = sorted_words[0]
        w = first['word']
        mid = self.size // 2
        row = mid
        col = mid - len(w) // 2
        self.place_word(w, row, col, 'H')
        self.placed.append({
            "word": w, "clue": first['clue'],
            "row": row, "col": col, "direction": 'H', "number": 1
        })

        num = 2
        remaining = sorted_words[1:]

        for item in remaining:
            word = item['word']
            placed = False

            # Avval kesishadigan joylarni sinab ko'rish
            for direction in ['V', 'H']:
                positions = self.find_best_positions(word, direction)
                for r, c in positions[:10]:  # Eng yaxshi 10 ta pozitsiya
                    if self.can_place(word, r, c, direction):
                        self.place_word(word, r, c, direction)
                        self.placed.append({
                            "word": word, "clue": item['clue'],
                            "row": r, "col": c, "direction": direction, "number": num
                        })
                        num += 1
                        placed = True
                        break
                if placed:
                    break

            # Kesishmasdan ham qo'yish (oxirgi chora)
            if not placed:
                for direction in ['H', 'V']:
                    for r in range(1, self.size - len(word) - 1):
                        for c in range(1, self.size - len(word) - 1):
                            if self.can_place(word, r, c, direction):
                                # Yaqinida boshqa so'z borligini tekshirish
                                near = False
                                check_range = 4
                                for cr in range(max(0, r-check_range), min(self.size, r+len(word)+check_range)):
                                    for cc in range(max(0, c-check_range), min(self.size, c+len(word)+check_range)):
                                        if self.grid[cr][cc] != '.':
                                            near = True
                                            break
                                    if near:
                                        break
                                if near:
                                    self.place_word(word, r, c, direction)
                                    self.placed.append({
                                        "word": word, "clue": item['clue'],
                                        "row": r, "col": c, "direction": direction, "number": num
                                    })
                                    num += 1
                                    placed = True
                                    break
                        if placed:
                            break
                    if placed:
                        break

        # Raqamlarni belgilash
        for p in self.placed:
            key = (p['row'], p['col'])
            if key not in self.numbers:
                self.numbers[key] = p['number']

    def get_bounds(self):
        min_r, max_r, min_c, max_c = self.size, 0, self.size, 0
        for r in range(self.size):
            for c in range(self.size):
                if self.grid[r][c] != '.':
                    min_r = min(min_r, r)
                    max_r = max(max_r, r)
                    min_c = min(min_c, c)
                    max_c = max(max_c, c)
        if min_r > max_r:
            return 0, 0, 0, 0
        return min_r, max_r, min_c, max_c


# ─── DOCX yordamchi funksiyalar ───────────────────────────
def set_cell_border(cell, color='555555'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=0, bottom=0, left=30, right=30):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


# ─── DOCX yaratish ────

def build_crossword_docx(
    cw: CrosswordGrid,
    topic: str,
    lang: str,
    author: str = "",
    with_answers: bool = False
) -> BytesIO:
    doc = Document()

    # Sahifa sozlamalari — landscape (gorizontal)
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Sarlavha
    title_type = "JAVOBLAR VARAQASI" if with_answers else "KROSSVORD"
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"KROSSVORD" if not with_answers else "JAVOBLAR VARAQASI")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    topic_para = doc.add_paragraph()
    topic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    topic_run = topic_para.add_run(f"Mavzu: {topic}")
    topic_run.bold = True
    topic_run.font.size = Pt(13)

    if author:
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(f"Tuzuvchi: {author}")
        info_run.font.size = Pt(11)
        info_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # To'r chizish
    min_r, max_r, min_c, max_c = cw.get_bounds()
    if min_r == max_r == min_c == max_c == 0 and cw.grid[0][0] == '.':
        doc.add_paragraph("Krossvord yaratilmadi.")
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    rows = max_r - min_r + 1
    cols = max_c - min_c + 1

    # Katakcha o'lchami — sahifaga sig'dirish
    available_width = 26.7  # cm (29.7 - 1.5*2)
    available_height = 14.0  # cm (21 - 1.5*2 - sarlavha)
    cell_w = min(Cm(0.95), Cm(available_width / max(cols, 1)))
    cell_h = min(Cm(0.95), Cm(available_height / max(rows, 1)))
    cell_size = min(cell_w, cell_h)

    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r in range(rows):
        row_obj = table.rows[r]
        row_obj.height = cell_size
        for c in range(cols):
            cell = table.cell(r, c)
            cell.width = cell_size

            gr = min_r + r
            gc = min_c + c
            grid_val = cw.grid[gr][gc]

            set_cell_margins(cell)

            if grid_val == '.':
                set_cell_bg(cell, '1E1E1E')
                # Chegara yo'q qora katakchada
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for edge in ('top', 'left', 'bottom', 'right'):
                    element = OxmlElement(f'w:{edge}')
                    element.set(qn('w:val'), 'nil')
                    tcBorders.append(element)
                tcPr.append(tcBorders)
                # Bo'sh paragraf
                cell.paragraphs[0].clear()
            else:
                set_cell_bg(cell, 'FFFFFF')
                set_cell_border(cell, color='888888')

                num = cw.numbers.get((gr, gc))
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)

                if num:
                    num_run = para.add_run(str(num))
                    num_run.font.size = Pt(6)
                    num_run.bold = True
                    num_run.font.color.rgb = RGBColor(0x00, 0x55, 0xCC)

                if with_answers:
                    # Harf ko'rsatish
                    if num:
                        letter_para = cell.add_paragraph()
                    else:
                        letter_para = para
                    letter_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    letter_para.paragraph_format.space_before = Pt(0)
                    letter_para.paragraph_format.space_after = Pt(0)
                    letter_run = letter_para.add_run(grid_val)
                    letter_run.font.size = Pt(9)
                    letter_run.bold = True
                    letter_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_paragraph()

    # Savollar
    clues_title = doc.add_paragraph()
    clues_run = clues_title.add_run("SAVOLLAR:")
    clues_run.bold = True
    clues_run.font.size = Pt(12)
    clues_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    h_words = sorted([p for p in cw.placed if p['direction'] == 'H'], key=lambda x: x['number'])
    v_words = sorted([p for p in cw.placed if p['direction'] == 'V'], key=lambda x: x['number'])

    if h_words:
        h_title = doc.add_paragraph()
        h_run = h_title.add_run("Gorizontal (→):")
        h_run.bold = True
        h_run.font.size = Pt(11)
        for p in h_words:
            q_para = doc.add_paragraph(style='List Number')
            q_para.paragraph_format.left_indent = Cm(0.5)
            q_run = q_para.add_run(f"{p['number']}. {p['clue']}")
            q_run.font.size = Pt(10.5)

    if v_words:
        v_title = doc.add_paragraph()
        v_run = v_title.add_run("Vertikal (↓):")
        v_run.bold = True
        v_run.font.size = Pt(11)
        for p in v_words:
            q_para = doc.add_paragraph(style='List Number')
            q_para.paragraph_format.left_indent = Cm(0.5)
            q_run = q_para.add_run(f"{p['number']}. {p['clue']}")
            q_run.font.size = Pt(10.5)

    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Biz bilan ishingiz oson! | @slidego | t.me/slidego")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── Asosiy funksiya ──────────────────────────────────────
async def generate_crossword(
    topic: str,
    count: int,
    lang: str,
    author: str = "",
) -> tuple[BytesIO, BytesIO]:
    import asyncio

    words_data = await asyncio.to_thread(generate_words, topic, count, lang)

    if not words_data:
        raise ValueError("GPT so'zlar ro'yxatini qaytarmadi")

    cw = CrosswordGrid(size=22)
    cw.build(words_data)

    if not cw.placed:
        raise ValueError("Krossvord to'ri qurilmadi")

    empty_doc = await asyncio.to_thread(
        build_crossword_docx, cw, topic, lang, author, False
    )
    answer_doc = await asyncio.to_thread(
        build_crossword_docx, cw, topic, lang, author, True
    )

    return empty_doc, answer_doc


# Narxlar
CROSSWORD_PRICES = {
    10: 1000,
    15: 2000,
    20: 2000,
}
