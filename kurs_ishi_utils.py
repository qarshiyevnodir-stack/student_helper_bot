"""
kurs_ishi_utils.py
Kurs ishi va Bitiruv malakaviy ishi (BMI) generatori.
Bitta mega-so'rov bilan barcha kontent olinadi, so'ng professional DOCX yaratiladi.
"""

import os
import io
import json
import logging
import requests
import tempfile
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from io import BytesIO
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)
client = OpenAI()

# ─────────────────────────────────────────────
# Konfiguratsiya
# ─────────────────────────────────────────────

WORK_TYPE_NAMES = {
    "kurs_ishi": "KURS ISHI",
    "bmi":       "BITIRUV MALAKAVIY ISHI",
}

LANGUAGE_NAMES = {
    "uz": "o'zbek",
    "ru": "rus",
    "en": "ingliz",
    "ko": "kores",
    "zh": "xitoy",
    "de": "nemis",
}

def get_chapter_count(page_count: int, work_type: str) -> int:
    """Sahifa soniga qarab bob sonini aniqlash."""
    if work_type == "bmi":
        if page_count <= 45:
            return 3
        elif page_count <= 60:
            return 4
        else:
            return 5
    else:  # kurs_ishi
        if page_count <= 25:
            return 2
        elif page_count <= 35:
            return 3
        else:
            return 3

def get_words_per_section(page_count: int, work_type: str) -> dict:
    """Har bir bo'lim uchun so'z soni."""
    # 1 sahifa ≈ 300-350 so'z (Times New Roman 14pt, 1.5 interval)
    if work_type == "bmi":
        if page_count <= 45:
            return {"kirish": 600, "bob_paragraf": 500, "xulosa": 500}
        elif page_count <= 60:
            return {"kirish": 700, "bob_paragraf": 600, "xulosa": 600}
        else:
            return {"kirish": 800, "bob_paragraf": 700, "xulosa": 700}
    else:
        if page_count <= 25:
            return {"kirish": 500, "bob_paragraf": 450, "xulosa": 400}
        elif page_count <= 35:
            return {"kirish": 600, "bob_paragraf": 500, "xulosa": 450}
        else:
            return {"kirish": 700, "bob_paragraf": 550, "xulosa": 500}


# ─────────────────────────────────────────────
# Mega-so'rov: barcha kontent bir vaqtda
# ─────────────────────────────────────────────

def generate_all_content(
    topic: str,
    language: str,
    work_type: str,
    page_count: int,
    subject: str = "",
) -> dict:
    """Bitta GPT so'rovida barcha kontent olish."""

    chapter_count = get_chapter_count(page_count, work_type)
    words = get_words_per_section(page_count, work_type)
    lang_name = LANGUAGE_NAMES.get(language, "o'zbek")
    work_name = WORK_TYPE_NAMES.get(work_type, "KURS ISHI")
    subject_hint = f" (fan: {subject})" if subject else ""

    # Bob nomlari uchun ko'rsatma
    if chapter_count == 2:
        chapters_desc = """
- "bob_1": {"nomi": "Nazariy asoslar", "paragraflar": ["1.1 ...", "1.2 ...", "1.3 ..."]}
- "bob_2": {"nomi": "Tahliliy qism", "paragraflar": ["2.1 ...", "2.2 ...", "2.3 ..."]}
"""
    elif chapter_count == 3:
        chapters_desc = """
- "bob_1": {"nomi": "Nazariy asoslar", "paragraflar": ["1.1 ...", "1.2 ...", "1.3 ..."]}
- "bob_2": {"nomi": "Tahliliy qism", "paragraflar": ["2.1 ...", "2.2 ...", "2.3 ..."]}
- "bob_3": {"nomi": "Tavsiyalar va yechimlar", "paragraflar": ["3.1 ...", "3.2 ...", "3.3 ..."]}
"""
    else:
        chapters_desc = """
- "bob_1": {"nomi": "Nazariy asoslar", "paragraflar": ["1.1 ...", "1.2 ...", "1.3 ..."]}
- "bob_2": {"nomi": "Tahliliy qism", "paragraflar": ["2.1 ...", "2.2 ...", "2.3 ..."]}
- "bob_3": {"nomi": "Tavsiyalar va yechimlar", "paragraflar": ["3.1 ...", "3.2 ...", "3.3 ..."]}
- "bob_4": {"nomi": "Amaliy tatbiq", "paragraflar": ["4.1 ...", "4.2 ...", "4.3 ..."]}
"""

    prompt = f"""Siz {lang_name} tilida professional {work_name} yozuvchi ekspertsiz.
Mavzu: "{topic}"{subject_hint}
Ish turi: {work_name}
Hajm: taxminan {page_count} sahifa

Quyidagi JSON formatda to'liq kontent yarating. Har bir matn KAMIDA ko'rsatilgan so'z sonida bo'lsin.

{{
  "kirish": {{
    "dolzarblik": "Mavzuning dolzarbligi haqida KAMIDA {words['kirish']//5} so'z",
    "maqsad": "Ishning maqsadi (2-3 gap)",
    "vazifalar": ["vazifa 1", "vazifa 2", "vazifa 3", "vazifa 4"],
    "obekt": "Tadqiqot ob'ekti (1-2 gap)",
    "predmet": "Tadqiqot predmeti (1-2 gap)",
    "metodlar": "Qo'llanilgan metodlar (2-3 gap)",
    "tuzilma": "Ishning tuzilmasi haqida qisqacha (2-3 gap)",
    "matn": "Kirish bo'limining to'liq matni KAMIDA {words['kirish']} so'z"
  }},
  "boblar": [
    {{
      "bob_nomi": "I BOB. [Bob nomi]",
      "paragraflar": [
        {{
          "sarlavha": "1.1 [Paragraf nomi]",
          "matn": "KAMIDA {words['bob_paragraf']} so'z. Ilmiy uslubda, faktlar va tahlil bilan."
        }},
        {{
          "sarlavha": "1.2 [Paragraf nomi]",
          "matn": "KAMIDA {words['bob_paragraf']} so'z.",
          "jadval": {{
            "sarlavha": "Jadval 1.1 — [Jadval nomi]",
            "ustunlar": ["Ko'rsatkich", "Qiymat", "Izoh"],
            "qatorlar": [
              ["Ko'rsatkich 1", "Qiymat 1", "Izoh 1"],
              ["Ko'rsatkich 2", "Qiymat 2", "Izoh 2"],
              ["Ko'rsatkich 3", "Qiymat 3", "Izoh 3"],
              ["Ko'rsatkich 4", "Qiymat 4", "Izoh 4"]
            ]
          }}
        }},
        {{
          "sarlavha": "1.3 [Paragraf nomi]",
          "matn": "KAMIDA {words['bob_paragraf']} so'z.",
          "iqtibos": "Muhim ta'rif yoki iqtibos (1-2 gap)"
        }}
      ],
      "bob_xulosasi": "Bob xulosasi (3-4 gap)"
    }}
    {"," if chapter_count >= 2 else ""}
    {{
      "bob_nomi": "II BOB. [Bob nomi]",
      "paragraflar": [
        {{
          "sarlavha": "2.1 [Paragraf nomi]",
          "matn": "KAMIDA {words['bob_paragraf']} so'z.",
          "grafik": {{
            "sarlavha": "Rasm 2.1 — [Grafik nomi]",
            "turi": "bar",
            "yorliqlar": ["Kategoriya 1", "Kategoriya 2", "Kategoriya 3", "Kategoriya 4", "Kategoriya 5"],
            "qiymatlar": [45, 72, 38, 85, 61],
            "x_label": "Kategoriyalar",
            "y_label": "Qiymatlar"
          }}
        }},
        {{
          "sarlavha": "2.2 [Paragraf nomi]",
          "matn": "KAMIDA {words['bob_paragraf']} so'z."
        }},
        {{
          "sarlavha": "2.3 [Paragraf nomi]",
          "matn": "KAMIDA {words['bob_paragraf']} so'z."
        }}
      ],
      "bob_xulosasi": "Bob xulosasi (3-4 gap)"
    }}
    {generate_extra_chapters_json(chapter_count, words)}
  ],
  "xulosa": {{
    "umumiy_xulosa": "KAMIDA {words['xulosa']} so'z. Asosiy natijalar va xulosalar.",
    "tavsiyalar": ["Tavsiya 1", "Tavsiya 2", "Tavsiya 3", "Tavsiya 4", "Tavsiya 5"]
  }},
  "adabiyotlar": [
    "1. [Muallif]. [Kitob nomi]. — [Shahar]: [Nashriyot], [Yil]. — [Sahifalar] b.",
    "2. ...",
    "3. ...",
    "... (jami {get_ref_count(page_count, work_type)} ta manba)"
  ],
  "unsplash_keyword": "mavzuga oid inglizcha 1-2 so'z (rasm qidirish uchun)"
}}

MUHIM: Barcha matnlar {lang_name} tilida bo'lsin. Ilmiy uslubda yozing. Har bir paragraf matnida kamida {words['bob_paragraf']} so'z bo'lsin."""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"},
        max_tokens=16000,
        temperature=0.7,
    )
    return json.loads(response.choices[0].message.content)


def generate_extra_chapters_json(chapter_count: int, words: dict) -> str:
    """3 va 4-bob uchun JSON qo'shimcha."""
    result = ""
    if chapter_count >= 3:
        result += f""",
    {{
      "bob_nomi": "III BOB. [Bob nomi]",
      "paragraflar": [
        {{
          "sarlavha": "3.1 [Paragraf nomi]",
          "matn": "KAMIDA {words['bob_paragraf']} so'z."
        }},
        {{
          "sarlavha": "3.2 [Paragraf nomi]",
          "matn": "KAMIDA {words['bob_paragraf']} so'z."
        }},
        {{
          "sarlavha": "3.3 [Paragraf nomi]",
          "matn": "KAMIDA {words['bob_paragraf']} so'z."
        }}
      ],
      "bob_xulosasi": "Bob xulosasi (3-4 gap)"
    }}"""
    if chapter_count >= 4:
        result += f""",
    {{
      "bob_nomi": "IV BOB. [Bob nomi]",
      "paragraflar": [
        {{
          "sarlavha": "4.1 [Paragraf nomi]",
          "matn": "KAMIDA {words['bob_paragraf']} so'z."
        }},
        {{
          "sarlavha": "4.2 [Paragraf nomi]",
          "matn": "KAMIDA {words['bob_paragraf']} so'z."
        }},
        {{
          "sarlavha": "4.3 [Paragraf nomi]",
          "matn": "KAMIDA {words['bob_paragraf']} so'z."
        }}
      ],
      "bob_xulosasi": "Bob xulosasi (3-4 gap)"
    }}"""
    return result


def get_ref_count(page_count: int, work_type: str) -> int:
    if work_type == "bmi":
        return 30 if page_count >= 60 else 25
    return 20 if page_count >= 35 else 15


# ─────────────────────────────────────────────
# Rasm olish
# ─────────────────────────────────────────────

def fetch_image(keyword: str) -> str | None:
    """Unsplash dan rasm URL olish, yo'qsa Picsum."""
    try:
        url = f"https://source.unsplash.com/800x400/?{keyword.replace(' ', ',')}"
        r = requests.get(url, timeout=10, allow_redirects=True)
        if r.status_code == 200 and "image" in r.headers.get("Content-Type", ""):
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
            tmp.write(r.content)
            tmp.close()
            return tmp.name
    except Exception:
        pass
    try:
        r = requests.get("https://picsum.photos/800/400", timeout=8)
        if r.status_code == 200:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
            tmp.write(r.content)
            tmp.close()
            return tmp.name
    except Exception:
        pass
    return None


def create_bar_chart(grafik: dict) -> str | None:
    """Bar grafik yaratib PNG fayl yo'lini qaytarish."""
    try:
        fig, ax = plt.subplots(figsize=(7, 4))
        yorliqlar = grafik.get("yorliqlar", [])
        qiymatlar = grafik.get("qiymatlar", [])
        colors = ["#1a3a6b", "#2e5fa3", "#4a7fc1", "#6b9fd4", "#8dbfe8"]
        ax.bar(yorliqlar, qiymatlar, color=colors[:len(yorliqlar)], edgecolor="white", linewidth=0.5)
        ax.set_xlabel(grafik.get("x_label", ""), fontsize=10)
        ax.set_ylabel(grafik.get("y_label", ""), fontsize=10)
        ax.set_title(grafik.get("sarlavha", ""), fontsize=11, fontweight="bold", pad=10)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.grid(axis="y", alpha=0.3)
        plt.xticks(rotation=15, ha="right", fontsize=9)
        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        plt.savefig(tmp.name, dpi=150, bbox_inches="tight")
        plt.close()
        return tmp.name
    except Exception as e:
        logger.error(f"Grafik yaratishda xatolik: {e}")
        return None


# ─────────────────────────────────────────────
# DOCX formatlash yordamchi funksiyalar
# ─────────────────────────────────────────────

def set_font(run, size=14, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    """Bob sarlavhasi qo'shish."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper() if level == 1 else text)
    set_font(run, size=14, bold=True)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_body_text(doc, text):
    """Asosiy matn qo'shish (1.5 interval, 1.25 sm chap chekinish)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = Pt(21)  # 1.5 interval
    p.paragraph_format.space_after = Pt(0)
    return p


def add_blockquote(doc, text):
    """Iqtibos bloki (chap chekinish, kursiv)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f'"{text}"')
    set_font(run, size=13, italic=True, color=(70, 70, 70))
    p.paragraph_format.left_indent = Cm(2)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Chap chiziq
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "6")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), "1a3a6b")
    pBdr.append(left)
    pPr.append(pBdr)
    return p


def add_table(doc, jadval: dict):
    """Professional jadval qo'shish."""
    sarlavha = jadval.get("sarlavha", "")
    ustunlar = jadval.get("ustunlar", [])
    qatorlar = jadval.get("qatorlar", [])
    if not ustunlar or not qatorlar:
        return

    # Jadval sarlavhasi
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(sarlavha)
    set_font(run, size=12, bold=True, italic=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1 + len(qatorlar), cols=len(ustunlar))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Sarlavha qatori
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(ustunlar):
        hdr_cells[i].text = col
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                set_font(run, size=12, bold=True)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Ko'k fon
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1a3a6b")
        tcPr.append(shd)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Ma'lumot qatorlari
    for row_idx, qator in enumerate(qatorlar):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(qator[:len(ustunlar)]):
            row_cells[col_idx].text = str(cell_text)
            for para in row_cells[col_idx].paragraphs:
                for run in para.runs:
                    set_font(run, size=12)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()


def add_image(doc, img_path: str, caption: str = "", width_cm: float = 14):
    """Rasm qo'shish."""
    if not img_path or not os.path.exists(img_path):
        return
    try:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
        if caption:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = p_cap.add_run(caption)
            set_font(run_cap, size=11, italic=True)
            p_cap.paragraph_format.space_after = Pt(12)
    except Exception as e:
        logger.error(f"Rasm qo'shishda xatolik: {e}")


def add_page_break(doc):
    doc.add_page_break()


# ─────────────────────────────────────────────
# Muqova sahifasi
# ─────────────────────────────────────────────

def add_page_border(doc):
    """Hujjatning 1-sahifasiga qora to'rtburchak ramka qo'shadi."""
    for section in doc.sections:
        sectPr = section._sectPr
        for old in sectPr.findall(qn('w:pgBorders')):
            sectPr.remove(old)
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        pgBorders.set(qn('w:display'), 'firstPage')
        for edge in ('top', 'left', 'bottom', 'right'):
            border_el = OxmlElement(f'w:{edge}')
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '18')
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), '000000')
            pgBorders.append(border_el)
        sectPr.append(pgBorders)


def build_title_page(doc, topic, work_type, name_surname, university, faculty, subject, teacher):
    """Rasmiy muqova sahifasi."""
    work_name = WORK_TYPE_NAMES.get(work_type, "KURS ISHI")

    # 1-sahifaga qora ramka qo'shish
    add_page_border(doc)

    # Universitet (shrift 14 + 6 = 20 pt)
    if university:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(university.upper())
        set_font(run, size=20, bold=True)
        p.paragraph_format.space_after = Pt(4)

    # Fakultet
    if faculty:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(faculty)
        set_font(run, size=13)
        p.paragraph_format.space_after = Pt(4)

    # Kafedra / Fan
    if subject:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'"{ subject}" fani bo\'yicha')
        set_font(run, size=13, italic=True)
        p.paragraph_format.space_after = Pt(6)

    # 3 ta bo'sh qator (ish turi oldidan)
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # Ish turi (shrift 16 + 8 = 24 pt)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(work_name)
    set_font(run, size=24, bold=True, color=(0, 0, 0))
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # 1 ta bo'sh qator (ish turi bilan mavzu orasida)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    # Mavzu
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Mavzu: "{topic}"')
    set_font(run, size=15, bold=True)
    p.paragraph_format.space_after = Pt(40)

    # Muallif va rahbar
    if name_surname or teacher:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if name_surname:
            run = p.add_run(f"Bajardi: {name_surname}\n")
            set_font(run, size=13)
        if teacher:
            run2 = p.add_run(f"Ilmiy rahbar: {teacher}")
            set_font(run2, size=13)
        p.paragraph_format.right_indent = Cm(2)
        p.paragraph_format.space_after = Pt(60)

    add_page_break(doc)


# ─────────────────────────────────────────────
# Mundarija
# ─────────────────────────────────────────────

def build_mundarija(doc, content: dict):
    """Mundarija sahifasi - 2-rasmdagi kabi to'liq kenglikda tab bilan."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement
    from docx.shared import Twips

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MUNDARIJA")
    set_font(run, size=14, bold=True)
    p.paragraph_format.space_after = Pt(16)

    # Sahifa kengligi: A4 (21cm) - chap chegara (3cm) - o'ng chegara (1.5cm) = 16.5cm = 9354 twips
    # Tab stop o'ng tomonga: 16.5cm = 9354 twips
    TAB_POS = Twips(9354)

    def add_toc_line(text, pg, bold=False, indent_cm=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if indent_cm > 0:
            p.paragraph_format.left_indent = Cm(indent_cm)

        # Tab stop sozlash (o'ng tomonga, nuqta to'ldiruvchi)
        pPr = p._p.get_or_add_pPr()
        tabs_el = _OxmlElement('w:tabs')
        tab_el = _OxmlElement('w:tab')
        tab_el.set(_qn('w:val'), 'right')
        tab_el.set(_qn('w:leader'), 'dot')
        tab_el.set(_qn('w:pos'), str(int(TAB_POS - Twips(indent_cm * 567))))
        tabs_el.append(tab_el)
        pPr.append(tabs_el)

        # Matn qo'shish
        run1 = p.add_run(text)
        set_font(run1, size=13, bold=bold)

        # Tab belgisi
        run_tab = p.add_run()
        tab_xml = _OxmlElement('w:tab')
        run_tab._r.append(tab_xml)

        # Sahifa raqami
        run2 = p.add_run(pg)
        set_font(run2, size=13, bold=bold)

        p.paragraph_format.space_after = Pt(2)
        return p

    items = [
        ("KIRISH", "3", True, 0),
    ]
    page_num = 4
    for i, bob in enumerate(content.get("boblar", []), 1):
        bob_nomi = bob.get("bob_nomi", f"BOB {i}")
        items.append((bob_nomi, str(page_num), True, 0))
        page_num += 2
        for para in bob.get("paragraflar", []):
            items.append((para.get('sarlavha', ''), str(page_num), False, 1.0))
            page_num += 1

    items.append(("XULOSA", str(page_num), True, 0))
    items.append(("FOYDALANILGAN ADABIYOTLAR", str(page_num + 2), True, 0))

    for text, pg, bold, indent in items:
        add_toc_line(text, pg, bold=bold, indent_cm=indent)

    add_page_break(doc)


# ─────────────────────────────────────────────
# Kirish bo'limi
# ─────────────────────────────────────────────

def build_kirish(doc, kirish: dict, img_path: str = None):
    """Kirish sahifasi."""
    add_heading(doc, "KIRISH", level=1)

    # Dolzarblik
    add_body_text(doc, kirish.get("dolzarblik", ""))

    # Asosiy kirish matni
    add_body_text(doc, kirish.get("matn", ""))

    # Maqsad va vazifalar
    p = doc.add_paragraph()
    run = p.add_run("Ishning maqsadi: ")
    set_font(run, size=14, bold=True)
    run2 = p.add_run(kirish.get("maqsad", ""))
    set_font(run2, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(8)

    # Vazifalar
    p = doc.add_paragraph()
    run = p.add_run("Ishning vazifalari:")
    set_font(run, size=14, bold=True)
    p.paragraph_format.first_line_indent = Cm(1.25)
    for i, vazifa in enumerate(kirish.get("vazifalar", []), 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {vazifa}")
        set_font(run, size=14)
        p.paragraph_format.left_indent = Cm(1.25)

    # Ob'ekt va predmet
    for label, key in [("Tadqiqot ob'ekti", "obekt"), ("Tadqiqot predmeti", "predmet"), ("Metodologiya", "metodlar")]:
        val = kirish.get(key, "")
        if val:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            set_font(run, size=14, bold=True)
            run2 = p.add_run(val)
            set_font(run2, size=14)
            p.paragraph_format.first_line_indent = Cm(1.25)

    # Rasm
    if img_path:
        doc.add_paragraph()
        add_image(doc, img_path, caption="Rasm 1 — Mavzuga oid ko'rgazmali material", width_cm=14)

    add_page_break(doc)


# ─────────────────────────────────────────────
# Bob bo'limlari
# ─────────────────────────────────────────────

def build_bob(doc, bob: dict, bob_idx: int):
    """Bir bob va uning paragraflarini qo'shish."""
    add_heading(doc, bob.get("bob_nomi", f"BOB {bob_idx}"), level=1)

    for para_idx, para in enumerate(bob.get("paragraflar", [])):
        # Paragraf sarlavhasi
        add_heading(doc, para.get("sarlavha", ""), level=2)

        # Asosiy matn
        matn = para.get("matn", "")
        if matn:
            add_body_text(doc, matn)

        # Jadval
        if "jadval" in para:
            add_table(doc, para["jadval"])

        # Grafik
        if "grafik" in para:
            chart_path = create_bar_chart(para["grafik"])
            if chart_path:
                add_image(doc, chart_path, caption=para["grafik"].get("sarlavha", ""), width_cm=13)

        # Iqtibos
        if "iqtibos" in para:
            add_blockquote(doc, para["iqtibos"])

    # Bob xulosasi
    bob_xulosa = bob.get("bob_xulosasi", "")
    if bob_xulosa:
        p = doc.add_paragraph()
        run = p.add_run("Bob bo'yicha xulosa: ")
        set_font(run, size=14, bold=True)
        run2 = p.add_run(bob_xulosa)
        set_font(run2, size=14)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_before = Pt(12)

    add_page_break(doc)


# ─────────────────────────────────────────────
# Xulosa
# ─────────────────────────────────────────────

def build_xulosa(doc, xulosa: dict):
    """Xulosa va tavsiyalar sahifasi."""
    add_heading(doc, "XULOSA VA TAVSIYALAR", level=1)

    add_body_text(doc, xulosa.get("umumiy_xulosa", ""))

    tavsiyalar = xulosa.get("tavsiyalar", [])
    if tavsiyalar:
        p = doc.add_paragraph()
        run = p.add_run("Asosiy tavsiyalar:")
        set_font(run, size=14, bold=True)
        p.paragraph_format.space_before = Pt(10)
        for i, tav in enumerate(tavsiyalar, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {tav}")
            set_font(run, size=14)
            p.paragraph_format.left_indent = Cm(1.25)

    add_page_break(doc)


# ─────────────────────────────────────────────
# Adabiyotlar
# ─────────────────────────────────────────────

def build_adabiyotlar(doc, adabiyotlar: list):
    """Foydalanilgan adabiyotlar sahifasi."""
    add_heading(doc, "FOYDALANILGAN ADABIYOTLAR", level=1)

    for ref in adabiyotlar:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        set_font(run, size=13)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(4)


# ─────────────────────────────────────────────
# Hujjat sahifa sozlamalari
# ─────────────────────────────────────────────

def setup_page(doc):
    """A4, GOST chegaralari: chap 3cm, o'ng 1.5cm, yuqori/past 2cm."""
    from docx.oxml.ns import qn
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)


# ─────────────────────────────────────────────
# Asosiy generator funksiya
# ─────────────────────────────────────────────

def generate_kurs_ishi(
    topic: str,
    language: str = "uz",
    work_type: str = "kurs_ishi",
    page_count: int = 25,
    name_surname: str = "",
    university: str = "",
    faculty: str = "",
    subject: str = "",
    teacher: str = "",
) -> BytesIO:
    """Kurs ishi yoki BMI yaratib BytesIO qaytaradi."""

    logger.info(f"Kurs ishi yaratilmoqda: {topic} | {work_type} | {page_count} sah | {language}")

    # 1. Kontent generatsiya
    content = generate_all_content(
        topic=topic,
        language=language,
        work_type=work_type,
        page_count=page_count,
        subject=subject,
    )

    # 2. Rasm olish
    keyword = content.get("unsplash_keyword", topic[:20])
    img_path = fetch_image(keyword)

    # 3. DOCX yaratish
    doc = Document()
    setup_page(doc)

    # Default font
    from docx.oxml.ns import qn as _qn
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    # 3.1 Muqova
    build_title_page(doc, topic, work_type, name_surname, university, faculty, subject, teacher)

    # 3.2 Mundarija
    build_mundarija(doc, content)

    # 3.3 Kirish
    build_kirish(doc, content.get("kirish", {}), img_path=img_path)

    # 3.4 Boblar
    for i, bob in enumerate(content.get("boblar", []), 1):
        build_bob(doc, bob, i)

    # 3.5 Xulosa
    build_xulosa(doc, content.get("xulosa", {}))

    # 3.6 Adabiyotlar
    build_adabiyotlar(doc, content.get("adabiyotlar", []))

    # 4. BytesIO ga saqlash
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    logger.info(f"Kurs ishi muvaffaqiyatli yaratildi: {topic}")
    return buf
