"""
tezis_utils.py — Tezis generatori
Ilmiy konferensiya, olimpiada, seminar va dissertatsiya tezislari uchun
professional DOCX hujjat yaratadi.
"""

import os
import json
import logging
import requests
import tempfile
from io import BytesIO
from openai import OpenAI

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

def get_client():
    """OpenAI client ni lazy yaratish."""
    return OpenAI()

# ─────────────────────────────────────────────
# Til sozlamalari
# ─────────────────────────────────────────────

LANG_LABELS = {
    "uz": "O'zbek",
    "ru": "Rus",
    "en": "Ingliz",
    "ko": "Kores",
    "zh": "Xitoy",
    "de": "Nemis",
}

LANG_PROMPTS = {
    "uz": "O'zbek tilida yoz.",
    "ru": "Пиши на русском языке.",
    "en": "Write in English.",
    "ko": "한국어로 작성하세요.",
    "zh": "用中文写。",
    "de": "Schreibe auf Deutsch.",
}

TEZIS_TYPES = {
    "konferensiya": "Konferensiya tezisi",
    "olimpiada": "Olimpiada tezisi",
    "seminar": "Seminar tezisi",
    "dissertatsiya": "Dissertatsiya tezisi",
}

# Tezis turi bo'yicha sahifa soni
# 1 A4 sahifa ≈ 350-400 so'z (Times New Roman 12pt, 1.5 interval)
TEZIS_PAGES = {
    1: {"sections": 2, "words_per_section": 250, "refs": 5},
    2: {"sections": 3, "words_per_section": 380, "refs": 6},
    3: {"sections": 3, "words_per_section": 500, "refs": 8},
    5: {"sections": 4, "words_per_section": 700, "refs": 12},
}


# ─────────────────────────────────────────────
# Mega-so'rov — barcha kontent 1 ta GPT so'rovda
# ─────────────────────────────────────────────

def generate_tezis_content(topic: str, tezis_type: str, lang: str, pages: int) -> dict:
    """Barcha tezis kontentini bitta GPT so'rovida oladi."""
    cfg = TEZIS_PAGES.get(pages, TEZIS_PAGES[2])
    lang_instruction = LANG_PROMPTS.get(lang, LANG_PROMPTS["uz"])
    type_label = TEZIS_TYPES.get(tezis_type, "Konferensiya tezisi")
    n_sections = cfg["sections"]
    words = cfg["words_per_section"]
    n_refs = cfg["refs"]

    section_names_map = {
        "uz": ["Muammo bayoni", "Tadqiqot metodologiyasi", "Asosiy natijalar", "Muhokama"],
        "ru": ["Постановка проблемы", "Методология исследования", "Основные результаты", "Обсуждение"],
        "en": ["Problem Statement", "Research Methodology", "Main Results", "Discussion"],
        "ko": ["문제 설명", "연구 방법론", "주요 결과", "토론"],
        "zh": ["问题陈述", "研究方法", "主要结果", "讨论"],
        "de": ["Problemstellung", "Forschungsmethodik", "Hauptergebnisse", "Diskussion"],
    }
    section_names = section_names_map.get(lang, section_names_map["uz"])[:n_sections]

    intro_label = {"uz": "Kirish", "ru": "Введение", "en": "Introduction",
                   "ko": "서론", "zh": "引言", "de": "Einleitung"}.get(lang, "Kirish")
    conclusion_label = {"uz": "Xulosa", "ru": "Заключение", "en": "Conclusion",
                        "ko": "결론", "zh": "结论", "de": "Schlussfolgerung"}.get(lang, "Xulosa")
    keywords_label = {"uz": "Kalit so'zlar", "ru": "Ключевые слова", "en": "Keywords",
                      "ko": "키워드", "zh": "关键词", "de": "Schlüsselwörter"}.get(lang, "Kalit so'zlar")
    refs_label = {"uz": "Adabiyotlar", "ru": "Литература", "en": "References",
                  "ko": "참고문헌", "zh": "参考文献", "de": "Literatur"}.get(lang, "Adabiyotlar")
    abstract_label = {"uz": "Annotatsiya", "ru": "Аннотация", "en": "Abstract",
                      "ko": "초록", "zh": "摘要", "de": "Zusammenfassung"}.get(lang, "Annotatsiya")

    sections_json = {f"section_{i+1}": f"{section_names[i]} bo'limi matni" for i in range(n_sections)}
    sections_schema = "\n".join([f'  "section_{i+1}_title": "{section_names[i]}",\n  "section_{i+1}_text": "KAMIDA {words} so\'z matn"' for i in range(n_sections)])

    intro_words = max(200, words // 2)
    conclusion_words = max(150, words // 2)
    abstract_words = 80 if pages <= 2 else 120

    prompt = f"""Sen ilmiy tezis yozuvchisisiz. {lang_instruction}

Mavzu: "{topic}"
Tezis turi: {type_label}
Sahifa soni: {pages} sahifa (1 sahifa = taxminan 400 so'z)

Quyidagi JSON formatida TO'LIQ va BATAFSIL tezis yarating:

{{
  "title": "Aniq va ilmiy sarlavha",
  "abstract": "KAMIDA {abstract_words} so'zlik annotatsiya — mavzu, maqsad, usul va natijalar",
  "keywords": ["kalit1", "kalit2", "kalit3", "kalit4", "kalit5"],
  "introduction": "KAMIDA {intro_words} so'z — mavzuning dolzarbligi, tadqiqot maqsadi, vazifalari, ob'ekti va predmeti",
{sections_schema},
  "conclusion": "KAMIDA {conclusion_words} so'z — asosiy xulosalar, amaliy ahamiyati va tavsiyalar",
  "references": [
    "1. Muallif A. (2023). Kitob nomi. Nashriyot.",
    "2. Muallif B. (2022). Maqola nomi. Jurnal nomi, 5(2), 10-25.",
    ... (jami {n_refs} ta manba)
  ],
  "unsplash_keyword": "mavzuga mos inglizcha 1-2 so'z (rasm uchun)"
}}

MUHIM TALABLAR:
- Har bir asosiy bo'lim KAMIDA {words} so'z bo'lsin — qisqa yozma!
- Kirish KAMIDA {intro_words} so'z bo'lsin
- Xulosa KAMIDA {conclusion_words} so'z bo'lsin
- Ilmiy uslubda, aniq faktlar va dalillar bilan yozing
- Har bir bo'limda kamida 2-3 xat boshi (paragraf) bo'lsin
- Adabiyotlar real mualliflar va nashrlar bilan ko'rinsin
- JSON formatidan chiqmang
- HECH QACHON qisqartirma yoki '...' ishlatmang, to'liq yozing"""

    # Jami kerakli so'z soni
    total_words_needed = pages * 400
    system_msg = (
        f"Sen ilmiy hujjat yozuvchisisiz. Sening vazifang {pages} sahifalik "
        f"to'liq ilmiy tezis yaratish. Jami hujjatda KAMIDA {total_words_needed} so'z bo'lishi SHART. "
        "Har bir bo'limni batafsil, aniq va to'liq yoz. Hech qachon qisqartirma ishlatma."
    )

    client = get_client()
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": prompt}
        ],
        response_format={"type": "json_object"},
        temperature=0.7,
        max_tokens=8000,
    )

    data = json.loads(response.choices[0].message.content)

    # Har bir bo'limni kengaytirish (agar so'z soni yetarli bo'lmasa)
    client2 = get_client()
    for i in range(1, n_sections + 1):
        key = f"section_{i}_text"
        text = data.get(key, "")
        word_count = len(text.split())
        if word_count < words * 0.8:  # 80% dan kam bo'lsa kengaytir
            expand_prompt = (
                f"{lang_instruction}\n\n"
                f"Quyidagi ilmiy bo'limni KAMIDA {words} so'zgacha kengaytir. "
                f"Mavzu: '{topic}'. Bo'lim nomi: '{section_names[i-1]}'.\n\n"
                f"Hozirgi matn:\n{text}\n\n"
                f"Kengaytirilgan matnni faqat JSON formatida qaytar: {{\"text\": \"...\"}}"
            )
            try:
                r2 = client2.chat.completions.create(
                    model="gpt-4.1-mini",
                    messages=[{"role": "user", "content": expand_prompt}],
                    response_format={"type": "json_object"},
                    temperature=0.7,
                    max_tokens=3000,
                )
                expanded = json.loads(r2.choices[0].message.content)
                if expanded.get("text") and len(expanded["text"].split()) > word_count:
                    data[key] = expanded["text"]
            except Exception as e:
                logger.warning(f"Bo'lim {i} kengaytirishda xatolik: {e}")

    return data


# ─────────────────────────────────────────────
# Yordamchi funksiyalar
# ─────────────────────────────────────────────

def fetch_image(keyword: str) -> str | None:
    """Unsplash yoki Picsum dan rasm yuklab oladi."""
    try:
        url = f"https://source.unsplash.com/800x400/?{keyword.replace(' ', ',')}"
        r = requests.get(url, timeout=10, allow_redirects=True)
        if r.status_code == 200 and 'image' in r.headers.get('Content-Type', ''):
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
            tmp.write(r.content)
            tmp.close()
            return tmp.name
    except Exception:
        pass
    try:
        r = requests.get("https://picsum.photos/800/400", timeout=8)
        if r.status_code == 200:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
            tmp.write(r.content)
            tmp.close()
            return tmp.name
    except Exception:
        pass
    return None


def set_cell_bg(cell, hex_color: str):
    """Jadval katakcha fon rangini o'rnatadi."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1, color_hex: str = "1B3A6B"):
    """Sarlavha qo'shadi."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    r, g, b = int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_paragraph(doc: Document, text: str, font_size: int = 12, justify: bool = True):
    """Paragraf qo'shadi."""
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def add_divider(doc: Document, color_hex: str = "1B3A6B"):
    """Ko'k ajratuvchi chiziq qo'shadi."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ─────────────────────────────────────────────
# DOCX yaratish
# ─────────────────────────────────────────────

def build_tezis_docx(
    content: dict,
    topic: str,
    tezis_type: str,
    lang: str,
    pages: int,
    author: str = "",
    institution: str = "",
    image_path: str = None,
) -> BytesIO:
    """Professional tezis DOCX hujjatini yaratadi."""

    doc = Document()
    cfg = TEZIS_PAGES.get(pages, TEZIS_PAGES[2])
    type_label = TEZIS_TYPES.get(tezis_type, "Konferensiya tezisi")
    lang_label = LANG_LABELS.get(lang, "O'zbek")

    # Sahifa sozlamalari
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Standart shrift
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # ─── 1-SAHIFA: Sarlavha ───
    # Tezis turi
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(type_label.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(27, 58, 107)

    # Sarlavha
    title = content.get("title", topic)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(27, 58, 107)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)

    # Ko'k chiziq
    add_divider(doc)

    # Muallif va muassasa
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(author)
        run.bold = True
        run.font.size = Pt(12)

    if institution:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(institution)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Annotatsiya
    abstract_labels = {
        "uz": "ANNOTATSIYA", "ru": "АННОТАЦИЯ", "en": "ABSTRACT",
        "ko": "초록", "zh": "摘要", "de": "ZUSAMMENFASSUNG"
    }
    abstract_label = abstract_labels.get(lang, "ANNOTATSIYA")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(abstract_label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(27, 58, 107)

    abstract = content.get("abstract", "")
    if abstract:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(abstract)
        run.font.size = Pt(11)
        run.italic = True
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.right_indent = Cm(1)

    # Kalit so'zlar
    keywords = content.get("keywords", [])
    if keywords:
        kw_labels = {
            "uz": "Kalit so'zlar", "ru": "Ключевые слова", "en": "Keywords",
            "ko": "키워드", "zh": "关键词", "de": "Schlüsselwörter"
        }
        kw_label = kw_labels.get(lang, "Kalit so'zlar")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"{kw_label}: ")
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(", ".join(keywords))
        run2.font.size = Pt(11)
        run2.italic = True
        p.paragraph_format.left_indent = Cm(1)

    add_divider(doc)
    doc.add_page_break()

    # ─── 2-SAHIFA: Kirish ───
    intro_labels = {
        "uz": "KIRISH", "ru": "ВВЕДЕНИЕ", "en": "INTRODUCTION",
        "ko": "서론", "zh": "引言", "de": "EINLEITUNG"
    }
    intro_label = intro_labels.get(lang, "KIRISH")
    add_heading(doc, intro_label)
    add_divider(doc)

    # Rasm (kirish sahifasida)
    if image_path:
        try:
            doc.add_picture(image_path, width=Cm(14))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            logger.warning(f"Rasm qo'shishda xatolik: {e}")

    intro = content.get("introduction", "")
    if intro:
        add_paragraph(doc, intro)

    doc.add_page_break()

    # ─── ASOSIY BO'LIMLAR ───
    n_sections = cfg["sections"]
    for i in range(1, n_sections + 1):
        title_key = f"section_{i}_title"
        text_key = f"section_{i}_text"
        sec_title = content.get(title_key, f"Bo'lim {i}")
        sec_text = content.get(text_key, "")

        add_heading(doc, f"{i}. {sec_title.upper()}")
        add_divider(doc)

        if sec_text:
            # Uzun matnni paragrafga bo'lish
            paragraphs = sec_text.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if para:
                    add_paragraph(doc, para)

        # Oxirgi bo'limdan oldin sahifa uzilishi
        if i < n_sections:
            doc.add_page_break()

    doc.add_page_break()

    # ─── XULOSA ───
    conclusion_labels = {
        "uz": "XULOSA", "ru": "ЗАКЛЮЧЕНИЕ", "en": "CONCLUSION",
        "ko": "결론", "zh": "结论", "de": "SCHLUSSFOLGERUNG"
    }
    conclusion_label = conclusion_labels.get(lang, "XULOSA")
    add_heading(doc, conclusion_label)
    add_divider(doc)

    conclusion = content.get("conclusion", "")
    if conclusion:
        add_paragraph(doc, conclusion)

    doc.add_page_break()

    # ─── ADABIYOTLAR ───
    refs_labels = {
        "uz": "FOYDALANILGAN ADABIYOTLAR", "ru": "СПИСОК ЛИТЕРАТУРЫ",
        "en": "REFERENCES", "ko": "참고문헌", "zh": "参考文献",
        "de": "LITERATURVERZEICHNIS"
    }
    refs_label = refs_labels.get(lang, "FOYDALANILGAN ADABIYOTLAR")
    add_heading(doc, refs_label)
    add_divider(doc)

    references = content.get("references", [])
    for ref in references:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(ref)
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)

    # BytesIO ga saqlash
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# Asosiy funksiya
# ─────────────────────────────────────────────

async def generate_tezis(
    topic: str,
    tezis_type: str,
    lang: str,
    pages: int,
    author: str = "",
    institution: str = "",
) -> BytesIO:
    """Tezis yaratadi va BytesIO qaytaradi."""
    import asyncio

    # 1. Kontent generatsiya (mega-so'rov)
    content = await asyncio.to_thread(
        generate_tezis_content, topic, tezis_type, lang, pages
    )

    # 2. Rasm olish
    keyword = content.get("unsplash_keyword", topic[:20])
    image_path = await asyncio.to_thread(fetch_image, keyword)

    # 3. DOCX yaratish
    docx_bytes = await asyncio.to_thread(
        build_tezis_docx,
        content, topic, tezis_type, lang, pages, author, institution, image_path
    )

    # Vaqtinchalik faylni o'chirish
    if image_path:
        try:
            os.unlink(image_path)
        except Exception:
            pass

    return docx_bytes
